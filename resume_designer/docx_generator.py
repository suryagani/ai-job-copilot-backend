from io import BytesIO

from docx import Document
from docx.shared import Pt


def build_docx(render_payload: dict) -> BytesIO:
    document = Document()
    typography = render_payload["layout"]["typography"]
    theme = render_payload["layout"]["theme"]

    normal_style = document.styles["Normal"]
    normal_style.font.name = theme["docx_font"]
    normal_style.font.size = Pt(typography["body_size"])

    for section in document.sections:
        margin_pts = render_payload["layout"]["spacing"]["margins"]
        section.top_margin = Pt(margin_pts)
        section.bottom_margin = Pt(margin_pts)
        section.left_margin = Pt(margin_pts)
        section.right_margin = Pt(margin_pts)

    title_p = document.add_paragraph()
    title_run = title_p.add_run(render_payload["title"])
    title_run.bold = True
    title_run.font.size = Pt(typography["title_size"])

    if render_payload["subtitle"]:
        subtitle_p = document.add_paragraph()
        subtitle_run = subtitle_p.add_run(render_payload["subtitle"])
        subtitle_run.italic = True
        subtitle_run.font.size = Pt(typography["subtitle_size"])

    for section in render_payload["sections"]:
        heading = document.add_paragraph()
        heading_run = heading.add_run(section["heading"])
        heading_run.bold = True
        heading_run.font.size = Pt(typography["heading_size"])

        for line in section["lines"]:
            if line.startswith("- "):
                document.add_paragraph(line[2:].strip(), style="List Bullet")
            else:
                document.add_paragraph(line)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer
