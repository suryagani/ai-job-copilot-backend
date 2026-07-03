from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field
from dotenv import load_dotenv
import os
import json
import smtplib
from io import BytesIO
from email.mime.text import MIMEText
from pathlib import Path
from json_repair import repair_json
from openai import OpenAI
from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import A4, letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.enums import TA_LEFT
from career_knowledge_engine import generate_career_knowledge
from cover_letter import generate_cover_letter_package
from linkedin_engine import generate_linkedin_optimization_package
from interview_engine import generate_interview_prep_package
from portfolio_engine import generate_portfolio_package
from resume_designer import render_resume_package
from resume_designer.regression_runner import run_regression_suite
from resume_models import select_resume_model
from resume_models.runtime import configure_runtime

# -----------------------
# Setup
# -----------------------
load_dotenv()
api_key = os.getenv("OPENAI_API_KEY")

if not api_key:
    raise RuntimeError("OPENAI_API_KEY not found. Put it in your .env file.")

client = OpenAI(api_key=api_key)
configure_runtime(client, lambda content: parse_json_response(content))

app = FastAPI(title="AI Job Copilot", version="1.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

DATA_DIR = Path(".")
JOB_ALERT_FILE = DATA_DIR / "job_alert_config.json"

SMTP_HOST = os.getenv("SMTP_HOST", "")
SMTP_PORT = int(os.getenv("SMTP_PORT", "587"))
SMTP_USER = os.getenv("SMTP_USER", "")
SMTP_PASSWORD = os.getenv("SMTP_PASSWORD", "")
SMTP_FROM = os.getenv("SMTP_FROM", SMTP_USER)

ROLE_CATALOG = [
    "Software Engineer",
    "Backend Developer",
    "Frontend Developer",
    "Full Stack Developer",
    "Mobile App Developer",
    "DevOps Engineer",
    "Cloud Engineer",
    "Site Reliability Engineer",
    "Platform Engineer",
    "Systems Administrator",
    "Network Engineer",
    "Cybersecurity Analyst",
    "IT Support Engineer",
    "QA Engineer",
    "QA Automation Engineer",
    "Automation Engineer",
    "RPA Developer",
    "Data Analyst",
    "Business Analyst",
    "Data Engineer",
    "Data Scientist",
    "BI Analyst",
    "AI / ML Engineer",
    "Machine Learning Engineer",
    "Generative AI Engineer",
    "MLOps Engineer",
    "Product Manager",
    "Project Manager",
    "Program Manager",
    "Scrum Master",
    "Operations Analyst",
    "Operations Executive",
    "Operations Manager",
    "Supply Chain Analyst",
    "Logistics Coordinator",
    "Procurement Specialist",
    "Sales Executive",
    "Business Development Executive",
    "Account Manager",
    "Marketing Executive",
    "Digital Marketing Specialist",
    "SEO Specialist",
    "Content Strategist",
    "Content Writer",
    "Social Media Manager",
    "Graphic Designer",
    "UI/UX Designer",
    "Motion Designer",
    "Recruiter",
    "Talent Acquisition Specialist",
    "HR Executive",
    "HR Generalist",
    "Customer Support Executive",
    "Customer Success Manager",
    "Accountant",
    "Financial Analyst",
    "Finance Executive",
    "Investment Analyst",
    "Audit Associate",
    "Tax Associate",
    "Banking Associate",
    "Administrative Assistant",
    "Executive Assistant",
    "Office Administrator",
    "Legal Associate",
    "Compliance Analyst",
    "Research Assistant",
    "Teacher",
    "Lecturer",
    "Instructional Designer",
    "Healthcare Administrator",
    "Pharmacist",
    "Nurse",
    "Medical Coder",
    "Clinical Data Analyst",
    "Electronics Engineer",
    "Embedded Systems Engineer",
    "VLSI Design Engineer",
    "Telecommunications Engineer",
    "ECE Engineer",
    "Electrical Engineer",
    "Power Systems Engineer",
    "Mechanical Engineer",
    "Automotive Engineer",
    "Manufacturing Engineer",
    "Production Engineer",
    "Industrial Engineer",
    "Civil Engineer",
    "Structural Engineer",
    "Site Engineer",
    "Construction Engineer",
    "Architect",
    "Interior Designer",
    "Environmental Engineer",
    "Chemical Engineer",
    "Process Engineer",
    "Food Technologist",
    "Agriculture Officer",
    "Geologist",
    "Quantity Surveyor",
]


# -----------------------
# Models
# -----------------------
class ProfileInput(BaseModel):
    target_role: str = Field(..., examples=["Python Developer"])
    target_location: str = Field("United Kingdom", examples=["United Kingdom"])
    headline: str = ""
    about: str
    experience: str = ""


class OptimizedProfile(BaseModel):
    headline: str
    about: str
    experience_bullets: list[str]
    top_keywords: list[str]


class LinkedInCopyInput(BaseModel):
    headline: str
    about: str
    experience_bullets: list[str]


class RoleSuggestionInput(BaseModel):
    about: str


class ScratchProfileInput(BaseModel):
    target_role: str
    education: str
    skills: str
    projects: str
    experience: str = ""
    career_goal: str
    target_location: str = "United Kingdom"


class ResumeOptimizerInput(BaseModel):
    target_role: str
    target_country: str = "United Kingdom"
    resume_text: str
    job_description: str = ""
    target_location: str = "United Kingdom"


class JobDescriptionRequest(BaseModel):
    job_title: str = ""
    company_name: str = ""
    country: str = ""
    industry: str = ""
    job_description: str


class AchievementRequest(BaseModel):
    target_role: str = ""
    target_country: str = "Global"
    experience_level: str = ""
    career_direction: str = ""
    work_experience: str = ""
    internships: str = ""
    projects: str = ""
    achievements: str = ""
    leadership_experience: str = ""
    technical_skills: str = ""
    transferable_skills: str = ""
    resume_text: str = ""


class ATSAnalysisRequest(BaseModel):
    target_role: str
    target_country: str = "Global"
    target_industry: str = ""
    experience_level: str = ""
    career_direction: str = ""
    resume_text: str = ""
    job_description: str = ""
    technical_skills: str = ""
    transferable_skills: str = ""
    tools_software: str = ""
    projects: str = ""
    work_experience: str = ""


class CareerKnowledgeRequest(BaseModel):
    target_role: str = ""
    target_industry: str = ""
    education: str = ""
    highest_qualification: str = ""
    current_background: str = ""
    career_direction: str = ""
    experience_level: str = ""


class ResumeIntelligenceRequest(BaseModel):
    target_role: str
    target_country: str = "Global"
    target_industry: str = ""
    career_direction: str = ""
    experience_level: str = ""
    current_background: str = ""
    highest_qualification: str = ""
    education_details: str = ""
    work_experience: str = ""
    internships: str = ""
    projects: str = ""
    technical_skills: str = ""
    transferable_skills: str = ""
    tools_software: str = ""
    certifications: str = ""
    achievements: str = ""
    leadership_experience: str = ""
    career_change: str = "No"
    current_field: str = ""
    target_field: str = ""


class ResumeIntelligenceInput(BaseModel):
    full_name: str = ""
    email: str = ""
    phone: str = ""
    location: str = ""
    linkedin_url: str = ""
    portfolio_url: str = ""
    target_role: str
    target_country: str = "Global"
    target_industry: str = ""
    career_direction: str = ""
    experience_level: str = ""
    current_background: str = ""
    highest_qualification: str = ""
    education_details: str = ""
    work_experience: str = ""
    internships: str = ""
    projects: str = ""
    technical_skills: str = ""
    transferable_skills: str = ""
    tools_software: str = ""
    certifications: str = ""
    achievements: str = ""
    leadership_experience: str = ""
    career_change: str = "No"
    current_field: str = ""
    target_field: str = ""
    preferred_resume_style: str = "Auto Recommend"
    job_description: str = ""
    company_name: str = ""


class LinkedInOptimizationInput(ResumeIntelligenceInput):
    resume_text: str = ""
    current_linkedin: str = ""
    current_headline: str = ""
    current_about: str = ""


class LinkedInOptimizationOutput(BaseModel):
    professional_headline: str
    about_section: str
    experience_rewrite: list[str]
    featured_section: list[str]
    skills_order: list[str]
    top_50_recruiter_keywords: list[str]
    headline_score: int
    linkedin_score: int
    recruiter_visibility_score: int
    creator_profile_suggestions: list[str]
    networking_suggestions: list[str]
    quality_notes: list[str]
    visibility_explanation: str
    linkedin_report_pdf_path: str
    linkedin_report_docx_path: str


class InterviewPrepInput(ResumeIntelligenceInput):
    resume_text: str = ""
    years_of_experience: str = ""


class InterviewPrepOutput(BaseModel):
    technical_questions: list[str]
    behavioral_questions: list[str]
    hr_questions: list[str]
    company_specific_questions: list[str]
    star_answer_examples: list[str]
    candidate_strengths: list[str]
    candidate_weaknesses: list[str]
    likely_follow_up_questions: list[str]
    interview_tips: list[str]
    confidence_score: int
    readiness_score: int
    mock_interview_plan: list[str]
    interview_report_pdf_path: str
    interview_report_docx_path: str


class PortfolioInput(ResumeIntelligenceInput):
    resume_text: str = ""
    years_of_experience: str = ""
    current_linkedin: str = ""
    current_headline: str = ""
    current_about: str = ""
    github_url: str = ""


class PortfolioOutput(BaseModel):
    professional_bio: str
    about_me: str
    personal_tagline: str
    project_showcase: list[str]
    project_case_studies: list[str]
    github_readme: str
    personal_website_content: str
    skills_section: list[str]
    timeline: list[str]
    contact_section: list[str]
    professional_footer: str
    seo_meta_title: str
    seo_meta_description: str
    selected_theme: str
    portfolio_score: int
    recruiter_score: int
    quality_notes: list[str]
    portfolio_html_path: str
    portfolio_readme_path: str
    portfolio_docx_path: str
    portfolio_pdf_path: str
    portfolio_json_path: str


class ResumeSkillGroup(BaseModel):
    category: str
    skills: list[str]


class ResumeIntelligenceAnalysisOutput(BaseModel):
    candidate_profile_type: str
    career_direction_detected: str
    career_change_detected: bool
    recommended_resume_model: str
    resume_length_rule: str
    target_market_strategy: str
    recruiter_positioning: str
    priority_sections: list[str]
    sections_to_minimize_or_remove: list[str]
    skill_grouping_strategy: list[ResumeSkillGroup]
    ats_keyword_strategy: list[str]
    missing_information: list[str]
    content_risk_flags: list[str]
    writing_guidance: str


class JobDescriptionAnalysisOutput(BaseModel):
    job_level: str
    career_direction: str
    industry: str
    experience_required: str
    education_expectation: str
    required_skills: list[str]
    preferred_skills: list[str]
    technical_skills: list[str]
    soft_skills: list[str]
    tools_and_platforms: list[str]
    ats_keywords: list[str]
    primary_responsibilities: list[str]
    leadership_required: bool
    communication_level: str
    problem_solving_level: str
    resume_focus_sections: list[str]
    recommended_resume_model: str
    important_certifications: list[str]
    company_type_guess: str
    hiring_priority_summary: str
    candidate_fit_strategy: str


class ExperienceAchievementOutput(BaseModel):
    original: str
    improved: str
    impact_type: str
    confidence: str


class ProjectAchievementOutput(BaseModel):
    project_name: str
    improved_bullets: list[str]
    technologies_detected: list[str]
    project_value: str


class AchievementIntelligenceOutput(BaseModel):
    experience_bullets: list[ExperienceAchievementOutput]
    project_bullets: list[ProjectAchievementOutput]
    leadership_bullets: list[str]
    transferable_achievement_bullets: list[str]
    missing_impact_questions: list[str]


class ATSKeywordPlacementOutput(BaseModel):
    summary: list[str]
    skills: list[str]
    experience: list[str]
    projects: list[str]


class ATSIntelligenceOutput(BaseModel):
    ats_score_estimate: int
    ats_readiness_level: str
    required_keywords: list[str]
    matching_keywords: list[str]
    missing_keywords: list[str]
    keyword_placement_strategy: ATSKeywordPlacementOutput
    formatting_risks: list[str]
    section_risks: list[str]
    keyword_stuffing_risk: str
    ats_improvement_actions: list[str]
    ats_strategy_note: str


class ResumePersonalizationOutput(BaseModel):
    tone: str
    writing_style: str
    resume_strategy: str
    priority_sections: list[str]
    de_emphasize_sections: list[str]
    industry_language: list[str]
    recommended_order: list[str]
    summary_strategy: str
    experience_strategy: str
    project_strategy: str
    skills_strategy: str
    certification_strategy: str
    overall_personalization_note: str


class CareerKnowledgeOutput(BaseModel):
    recommended_roles: list[str]
    recommended_industries: list[str]
    recommended_certifications: list[str]
    recommended_projects: list[str]
    recommended_resume_model: str
    career_transition_options: list[str]
    future_growth_roles: list[str]
    salary_progression_note: str


class ResumeBuildOutput(BaseModel):
    recommended_resume_style: str
    recommendation_reason: str
    professional_title: str
    executive_summary: str
    resume_length_rule: str
    target_market_strategy: str
    recruiter_positioning: str
    full_resume: str
    ats_keywords: list[str]
    skill_groups: list[ResumeSkillGroup]
    strengths: list[str]
    missing_information: list[str]
    improvement_suggestions: list[str]
    writing_quality_score: str
    resume_readability: str
    ats_readiness: str
    resume_confidence: str
    quality_score: int
    ats_readiness_score: int
    recruiter_readability_score: int
    role_alignment_score: int
    quality_issues_found: list[str]
    quality_fixes_applied: list[str]
    interview_probability: int
    recruiter_confidence: int
    first_impression: str
    shortlisting_decision: str
    top_strengths: list[str]
    top_concerns: list[str]
    missing_high_value_information: list[str]
    recommended_improvements: list[str]
    industry_keywords_missing: list[str]
    resume_competitiveness: str
    skill_match_percentage: int
    strong_matching_skills: list[str]
    missing_required_skills: list[str]
    missing_preferred_skills: list[str]
    resume_alignment_strategy: str
    ats_score_estimate: int
    ats_readiness_level: str
    matching_keywords: list[str]
    missing_keywords: list[str]
    ats_improvement_actions: list[str]
    personalization_score: int
    personalization_strategy: str
    personalization_notes: list[str]
    resume_pdf_path: str
    resume_docx_path: str
    selected_theme: str
    page_count: int
    render_quality_score: int


class ResumeOptimizerOutput(BaseModel):
    recommended_resume_style: str
    recommendation_reason: str
    optimized_resume: str
    ats_keywords: list[str]
    strengths: list[str]
    weaknesses_found: list[str]
    improvement_suggestions: list[str]
    ats_score_estimate: int
    ats_readiness_level: str
    matching_keywords: list[str]
    missing_keywords: list[str]
    ats_improvement_actions: list[str]
    personalization_score: int
    personalization_strategy: str
    personalization_notes: list[str]
    resume_pdf_path: str
    resume_docx_path: str
    selected_theme: str
    page_count: int
    render_quality_score: int
    interview_probability: int
    recruiter_confidence: int
    first_impression: str
    shortlisting_decision: str
    top_strengths: list[str]
    top_concerns: list[str]
    missing_high_value_information: list[str]
    recommended_improvements: list[str]
    industry_keywords_missing: list[str]
    resume_competitiveness: str
    skill_match_percentage: int
    strong_matching_skills: list[str]
    missing_required_skills: list[str]
    missing_preferred_skills: list[str]
    resume_alignment_strategy: str


class CoverLetterInput(ResumeIntelligenceInput):
    hiring_manager: str = ""
    tone: str = "Professional"
    years_of_experience: str = ""
    resume_text: str = ""


class CoverLetterOutput(BaseModel):
    cover_letter_text: str
    cover_letter_pdf_path: str
    cover_letter_docx_path: str
    cover_letter_quality_score: int
    ats_alignment_score: int
    recruiter_confidence: int


class ResumeReviewerInput(BaseModel):
    resume_text: str
    target_role: str = ""
    target_country: str = "Global ATS"
    job_description: str = ""


class ResumeReviewerOutput(BaseModel):
    ats_score_estimate: str
    score_reason: str
    missing_keywords: list[str]
    weak_sections: list[str]
    strong_sections: list[str]
    recruiter_feedback: str
    improvement_suggestions: list[str]


class ResumeExportSection(BaseModel):
    heading: str
    body: str


class ResumeExportInput(BaseModel):
    full_name: str = ""
    target_role: str = ""
    target_country: str = "United Kingdom"
    sections: list[ResumeExportSection]
    selected_theme: str = ""
    experience_level: str = ""


class HiringMessageInput(BaseModel):
    target_role: str
    company_name: str
    hiring_manager_name: str = ""
    job_context: str = ""
    personal_background: str
    target_location: str = "United Kingdom"


class HiringMessageOutput(BaseModel):
    connection_message: str
    outreach_message: str
    follow_up_message: str


class JobAlertInput(BaseModel):
    email: str
    target_role: str
    country: str
    city: str
    experience_level: str
    keywords: str = ""
    preferred_time: str = "09:00"


class JobAlertTestOutput(BaseModel):
    status: str
    message: str
    preview_subject: str
    preview_body: str


# -----------------------
# Helpers
# -----------------------
def parse_json_response(content: str) -> dict:
    try:
        return json.loads(content)
    except Exception:
        try:
            fixed = repair_json(content)
            return json.loads(fixed)
        except Exception as e:
            raise HTTPException(
                status_code=500,
                detail=f"Model returned invalid JSON. Error: {e}. Raw output: {content[:1000]}",
            )


def save_json_file(path: Path, data: dict) -> None:
    path.write_text(json.dumps(data, indent=2), encoding="utf-8")


def load_json_file(path: Path) -> dict:
    if not path.exists():
        return {}
    return json.loads(path.read_text(encoding="utf-8"))


def get_career_knowledge_context(candidate_data, **overrides) -> dict:
    profile = {
        "target_role": overrides.get("target_role", getattr(candidate_data, "target_role", "")),
        "target_industry": overrides.get("target_industry", getattr(candidate_data, "target_industry", "")),
        "education": overrides.get(
            "education",
            getattr(candidate_data, "education", "") or getattr(candidate_data, "education_details", ""),
        ),
        "highest_qualification": overrides.get("highest_qualification", getattr(candidate_data, "highest_qualification", "")),
        "current_background": overrides.get(
            "current_background",
            getattr(candidate_data, "current_background", "") or getattr(candidate_data, "about", "") or getattr(candidate_data, "resume_text", ""),
        ),
        "career_direction": overrides.get("career_direction", getattr(candidate_data, "career_direction", "")),
        "experience_level": overrides.get("experience_level", getattr(candidate_data, "experience_level", "")),
    }
    return generate_career_knowledge(profile)


def resume_text_to_sections(resume_text: str) -> list[ResumeExportSection]:
    known_headings = {
        "professional summary",
        "executive summary",
        "summary",
        "education",
        "projects",
        "project experience",
        "technical skills",
        "core skills",
        "core competencies",
        "skills",
        "experience",
        "professional experience",
        "relevant experience",
        "internships",
        "certifications",
        "leadership & achievements",
        "leadership and achievements",
        "achievements",
        "career highlights",
        "leadership competencies",
        "strategic achievements",
    }
    lines = [line.rstrip() for line in str(resume_text or "").splitlines()]
    if not lines:
        return []

    sections: list[ResumeExportSection] = []
    current_heading = "Professional Summary"
    current_lines: list[str] = []
    started = False

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        normalized = stripped.lower()
        if normalized in known_headings or (stripped.isupper() and len(stripped.split()) <= 4):
            if current_lines:
                sections.append(ResumeExportSection(heading=current_heading.title(), body="\n".join(current_lines)))
            current_heading = stripped.title()
            current_lines = []
            started = True
            continue
        if not started and "|" in stripped:
            continue
        if not started and stripped == lines[0].strip():
            continue
        current_lines.append(stripped)
        started = True

    if current_lines:
        sections.append(ResumeExportSection(heading=current_heading.title(), body="\n".join(current_lines)))

    return sections


def build_export_input_from_resume_response(response: dict, candidate_data, resume_key: str = "full_resume") -> ResumeExportInput:
    return ResumeExportInput(
        full_name=str(getattr(candidate_data, "full_name", "")).strip(),
        target_role=str(getattr(candidate_data, "target_role", "")).strip(),
        target_country=str(getattr(candidate_data, "target_country", "United Kingdom")).strip(),
        sections=resume_text_to_sections(response.get(resume_key, "")),
        selected_theme=str(response.get("selected_theme", "")).strip(),
        experience_level=str(getattr(candidate_data, "experience_level", "")).strip(),
    )


def get_country_template_settings(country: str) -> dict:
    normalized = (country or "").strip().lower()
    default_settings = {
        "page_size": A4,
        "summary_heading": "Professional Summary",
        "skills_heading": "Key Skills",
    }

    if "united states" in normalized or "canada" in normalized:
        return {
            "page_size": letter,
            "summary_heading": "Professional Summary",
            "skills_heading": "Core Skills",
        }

    if "australia" in normalized:
        return {
            "page_size": A4,
            "summary_heading": "Career Summary",
            "skills_heading": "Key Skills",
        }

    if "germany" in normalized:
        return {
            "page_size": A4,
            "summary_heading": "Professional Profile",
            "skills_heading": "Core Competencies",
        }

    return default_settings


def get_country_rules(country: str) -> str:
    normalized = (country or "").strip().lower()

    if "united states" in normalized or normalized == "usa":
        return (
            "Use a concise, impact-first US style: sharp professional summary, strong action verbs, "
            "keyword-dense skills, and a one-page preference unless seniority clearly justifies more."
        )
    if "canada" in normalized:
        return (
            "Use a clean Canadian professional style: strong summary, direct accomplishments, ATS-safe headings, "
            "and clear relevance to the target role."
        )
    if "united kingdom" in normalized or normalized == "uk":
        return (
            "Use a UK recruiter-friendly style: polished summary, clear responsibilities translated into results, "
            "and practical, professional wording without inflated claims."
        )
    if "australia" in normalized:
        return (
            "Use an Australia-ready style: practical professional summary, results-oriented bullet points, "
            "and strong relevance to the local job market."
        )
    if "germany" in normalized:
        return (
            "Use a structured Germany-aware style: precise wording, orderly sectioning, and strong emphasis on qualifications, "
            "technical credibility, and role alignment."
        )
    if "uae" in normalized or "united arab emirates" in normalized:
        return (
            "Use a UAE recruiter-friendly style: leadership, coordination, execution, and business impact should be clear, "
            "with polished and professional language suited to international employers."
        )

    return (
        "Use a global ATS-safe style: clean headings, role-first positioning, strong recruiter readability, "
        "and broad market-ready language without regional template labels."
    )


def get_resume_length_rules(experience_level: str) -> str:
    normalized = (experience_level or "").strip().lower()
    if normalized in {"student", "fresher"}:
        return "Maximum 1 page. Prioritize education, internships, projects, and the strongest role-relevant skills."
    if normalized in {"1–3 years", "1-3 years", "3–5 years", "3-5 years"}:
        return "Keep it to 1 page. Prioritize the most relevant experience, impact, and skills."
    if normalized in {"5–10 years", "5-10 years"}:
        return "Target 1 to 2 pages maximum. Keep only relevant experience and avoid repetition."
    if normalized in {"10+ years", "10+ yrs", "10+ year"}:
        return "Maximum 2 pages. Focus on leadership, strategic impact, and only the most relevant earlier experience."
    return "Keep the resume concise, relevant, and ATS-friendly with no repetition."


def get_resume_style_models() -> str:
    return (
        "Allowed resume styles:\n"
        "- Graduate ATS Resume\n"
        "- Technical Professional Resume\n"
        "- Career Switcher Resume\n"
        "- Business Professional Resume\n"
        "- Executive Leadership Resume"
    )


def clean_string_list(values) -> list[str]:
    if not isinstance(values, list):
        return []
    return [str(value).strip() for value in values if str(value).strip()]


def clean_skill_groups(values) -> list[dict]:
    if not isinstance(values, list):
        return []

    cleaned_groups = []
    for item in values:
        if not isinstance(item, dict):
            continue
        category = str(item.get("category", "")).strip()
        skills = clean_string_list(item.get("skills", []))
        if category and skills:
            cleaned_groups.append({"category": category, "skills": skills})

    return cleaned_groups


def sanitize_resume_text(resume_text: str) -> str:
    banned_markers = [
        "[No formal work experience]",
        "[Details not provided]",
        "[No certifications provided]",
        "[Year not provided]",
        "[Dates Not Provided]",
        "[Confidential Employer]",
        "[Institution Name Not Provided]",
    ]

    cleaned_lines = []
    for raw_line in str(resume_text or "").splitlines():
        line = raw_line.strip()
        if not line:
            if cleaned_lines and cleaned_lines[-1] != "":
                cleaned_lines.append("")
            continue

        normalized_line = line.lower()
        if any(marker.lower() in normalized_line for marker in banned_markers):
            continue
        if normalized_line.startswith("[") and normalized_line.endswith("]"):
            continue
        if "not provided" in normalized_line or "confidential" in normalized_line:
            continue

        cleaned_lines.append(raw_line.rstrip())

    while cleaned_lines and cleaned_lines[-1] == "":
        cleaned_lines.pop()

    return "\n".join(cleaned_lines)


def split_skills_text(*values: str) -> list[str]:
    raw_text = "\n".join(str(value or "") for value in values)
    for separator in [",", ";", "|", "/"]:
        raw_text = raw_text.replace(separator, "\n")

    cleaned = []
    seen = set()
    for item in raw_text.splitlines():
        skill = " ".join(item.strip().split())
        if not skill:
            continue
        key = skill.lower()
        if key in seen:
            continue
        seen.add(key)
        cleaned.append(skill)
    return cleaned


def title_case_skill(skill: str) -> str:
    value = str(skill or "").strip()
    if not value:
        return ""
    acronym_map = {
        "aws": "AWS",
        "sql": "SQL",
        "crm": "CRM",
        "erp": "ERP",
        "fpga": "FPGA",
        "vlsi": "VLSI",
        "cmos": "CMOS",
        "rtl": "RTL",
        "cadence": "Cadence",
        "synopsys": "Synopsys",
        "matlab": "MATLAB",
        "git": "Git",
        "linux": "Linux",
    }
    return acronym_map.get(value.lower(), value)


def get_role_skill_targets(target_role: str, career_direction: str, target_industry: str) -> dict:
    role_text = " ".join([target_role or "", career_direction or "", target_industry or ""]).lower()

    if any(keyword in role_text for keyword in ["devops", "sre", "cloud", "platform"]):
        return {
            "categories": [
                ("Cloud & Infrastructure", ["aws", "azure", "gcp", "docker", "kubernetes", "terraform", "ansible"]),
                ("CI/CD & Automation", ["jenkins", "github actions", "gitlab ci", "bash", "powershell", "git"]),
                ("Programming Languages", ["python", "java", "go", "javascript"]),
                ("Databases", ["sql", "mysql", "postgresql", "mongodb"]),
                ("Operating Systems", ["linux", "unix", "windows server"]),
            ],
            "missing": ["Terraform", "Kubernetes", "CI/CD pipelines", "Monitoring"],
        }
    if any(keyword in role_text for keyword in ["business analyst", "analyst", "reporting"]):
        return {
            "categories": [
                ("Business & Productivity Tools", ["excel", "powerpoint", "google sheets", "power bi", "tableau"]),
                ("Data & Analysis", ["sql", "python", "data analysis", "reporting", "dashboarding"]),
                ("Communication & Documentation", ["communication", "documentation", "presentation", "stakeholder management"]),
                ("Process & Coordination", ["coordination", "requirements gathering", "process improvement"]),
            ],
            "missing": ["Requirements gathering", "Stakeholder management", "Dashboarding"],
        }
    if any(keyword in role_text for keyword in ["vlsi", "embedded", "electronics", "ece"]):
        return {
            "categories": [
                ("VLSI Tools", ["xilinx ise", "modelsim", "cadence", "synopsys", "matlab"]),
                ("Design Flow", ["verilog", "vhdl", "vlsi basics", "cmos", "rtl", "timing analysis", "fpga"]),
                ("Operating Systems", ["linux", "unix", "windows"]),
                ("Core Electronics", ["digital electronics", "embedded c", "embedded systems", "circuit design"]),
                ("Professional Skills", ["problem solving", "analytical thinking", "teamwork", "communication"]),
            ],
            "missing": ["RTL Design", "Timing Analysis", "Cadence", "Synopsys"],
        }
    if any(keyword in role_text for keyword in ["hr", "human resources", "recruit", "talent"]):
        return {
            "categories": [
                ("Business Tools", ["excel", "powerpoint", "google sheets", "documentation", "presentation"]),
                ("Communication & Stakeholder Management", ["communication", "coordination", "teamwork", "stakeholder management"]),
                ("Operations & Coordination", ["scheduling", "onboarding", "event coordination", "coordination"]),
                ("Reporting & Documentation", ["documentation", "reporting", "presentation"]),
            ],
            "missing": ["Employee onboarding", "HR operations", "Recruitment coordination"],
        }
    if any(keyword in role_text for keyword in ["operations manager", "operations", "logistics", "supply chain"]):
        return {
            "categories": [
                ("Operations Management", ["staff scheduling", "inventory coordination", "vendor management", "shift leadership", "service quality management", "process improvement"]),
                ("Leadership & Team Management", ["team supervision", "staff training", "customer handling", "team coordination"]),
                ("Reporting & Business Tools", ["excel", "reporting", "pos systems", "inventory tracking systems"]),
                ("Customer & Service Delivery", ["customer issue resolution", "service quality", "operations coordination"]),
            ],
            "missing": ["Cost control", "Budget oversight", "KPI reporting"],
        }

    if (career_direction or "").strip().lower() == "technical":
        return {
            "categories": [
                ("Technical Skills", ["python", "java", "javascript", "sql", "aws", "docker", "linux"]),
                ("Tools & Platforms", ["git", "excel", "power bi", "tableau", "matlab"]),
                ("Professional Skills", ["problem solving", "analytical thinking", "teamwork", "communication"]),
            ],
            "missing": [],
        }

    return {
        "categories": [
            ("Business Tools", ["excel", "powerpoint", "google sheets", "crm", "erp"]),
            ("Communication & Documentation", ["communication", "documentation", "presentation", "reporting"]),
            ("Operations & Coordination", ["coordination", "scheduling", "stakeholder management", "teamwork"]),
        ],
        "missing": [],
    }


def generate_skill_intelligence(data, intelligence=None):
    role_targets = get_role_skill_targets(
        getattr(data, "target_role", ""),
        getattr(data, "career_direction", "") or (intelligence or {}).get("career_direction_detected", ""),
        getattr(data, "target_industry", ""),
    )
    all_skills = split_skills_text(
        getattr(data, "technical_skills", ""),
        getattr(data, "transferable_skills", ""),
        getattr(data, "tools_software", ""),
    )
    skills_map = {skill.lower(): skill for skill in all_skills}
    used = set()
    skill_groups = []

    for category, keywords in role_targets["categories"]:
        category_skills = []
        for keyword in keywords:
            for skill_lower, original in skills_map.items():
                if skill_lower in used:
                    continue
                if keyword in skill_lower or skill_lower in keyword:
                    category_skills.append(title_case_skill(original))
                    used.add(skill_lower)
                    break
        category_skills = clean_string_list(category_skills)[:8]
        if category_skills:
            skill_groups.append({"category": category, "skills": category_skills})

    remaining = [title_case_skill(original) for skill_lower, original in skills_map.items() if skill_lower not in used]
    remaining = clean_string_list(remaining)

    career_change = str(getattr(data, "career_change", "No")).strip().lower() in {"yes", "true", "1"}
    direction = ((intelligence or {}).get("career_direction_detected") or getattr(data, "career_direction", "") or "").strip().lower()
    if remaining:
        fallback_category = "Transferable Skills" if career_change or direction in {"non-technical", "management", "operations", "sales", "marketing", "finance", "customer support"} else "Additional Tools & Skills"
        skill_groups.append({"category": fallback_category, "skills": remaining[:8]})

    skill_groups = skill_groups[:6]
    priority_skills = []
    for group in skill_groups[:3]:
        priority_skills.extend(group["skills"][:3])
    priority_skills = clean_string_list(priority_skills)[:8]

    missing_role_skills = []
    for skill in role_targets.get("missing", []):
        if skill.lower() not in skills_map:
            missing_role_skills.append(skill)
    missing_role_skills = clean_string_list(missing_role_skills)[:6]

    if career_change:
        skill_positioning_note = "Emphasize transferable skills and business-ready tools first, and keep unrelated technical details secondary unless they directly support the target role."
    elif direction == "technical":
        skill_positioning_note = "Lead with role-critical technical capabilities, then supporting tools, operating environments, and professional strengths."
    else:
        skill_positioning_note = "Prioritize business tools, coordination, communication, reporting, and role-relevant operational strengths ahead of less relevant technical detail."

    return {
        "skill_groups": clean_skill_groups(skill_groups),
        "priority_skills": priority_skills,
        "missing_role_skills": missing_role_skills,
        "skill_positioning_note": skill_positioning_note,
    }


def detect_resume_length_label(experience_level: str) -> str:
    normalized = (experience_level or "").strip().lower()
    if normalized in {"student", "fresher", "1???3 years", "1-3 years", "3???5 years", "3-5 years"}:
        return "One Page"
    if normalized in {"5???10 years", "5-10 years"}:
        return "One to Two Pages"
    if normalized in {"10+ years", "10+ yrs", "10+ year"}:
        return "Two Pages"
    return "One Page"


def detect_career_direction(data: ResumeIntelligenceRequest) -> str:
    preferred = (data.career_direction or "").strip()
    if preferred:
        return preferred

    role_text = " ".join(
        [
            data.target_role,
            data.target_industry,
            data.current_background,
            data.target_field,
        ]
    ).lower()

    keyword_map = [
        ("Management", ["manager", "lead", "director", "head"]),
        ("Sales", ["sales", "account executive", "business development"]),
        ("Marketing", ["marketing", "seo", "content", "brand"]),
        ("Operations", ["operations", "logistics", "supply chain"]),
        ("Finance", ["finance", "account", "audit", "banking"]),
        ("Healthcare", ["healthcare", "nurse", "clinical", "medical"]),
        ("Customer Support", ["customer support", "customer success", "service desk"]),
        ("Technical", ["engineer", "developer", "data", "vlsi", "embedded", "cloud", "software", "it"]),
    ]

    for label, keywords in keyword_map:
        if any(keyword in role_text for keyword in keywords):
            return label

    return "Other"


def get_resume_intelligence_strategy(data: ResumeIntelligenceRequest) -> dict:
    experience = (data.experience_level or "").strip().lower()
    role_text = " ".join(
        [
            data.target_role,
            data.current_background,
            data.work_experience,
            data.leadership_experience,
            data.target_field,
        ]
    ).lower()
    career_change = (data.career_change or "").strip().lower() in {"yes", "true", "1"}
    direction = detect_career_direction(data)

    executive_signal = (
        experience in {"10+ years", "10+ yrs", "10+ year"}
        and any(keyword in role_text for keyword in ["manager", "head", "director", "lead", "operations manager", "restaurant manager"])
    )

    if executive_signal:
        profile_type = "Executive"
        resume_model = "Executive Leadership Resume"
    elif experience in {"10+ years", "10+ yrs", "10+ year", "5???10 years", "5-10 years"}:
        profile_type = "Senior"
        resume_model = "Business Professional Resume" if direction != "Technical" else "Technical Professional Resume"
    elif experience in {"3???5 years", "3-5 years", "1???3 years", "1-3 years"}:
        profile_type = "Early Career" if experience in {"1???3 years", "1-3 years"} else "Mid-Level"
        resume_model = "Technical Professional Resume" if direction == "Technical" else "Business Professional Resume"
    elif experience == "student":
        profile_type = "Student"
        resume_model = "Graduate ATS Resume"
    else:
        profile_type = "Fresher"
        resume_model = "Graduate ATS Resume" if direction == "Technical" else "Business Professional Resume"

    if career_change:
        resume_model = "Career Switcher Resume"

    return {
        "candidate_profile_type": profile_type,
        "career_direction_detected": direction,
        "career_change_detected": career_change,
        "recommended_resume_model": resume_model,
        "resume_length_rule": detect_resume_length_label(data.experience_level),
    }


def normalize_resume_intelligence_result(parsed: dict, data: ResumeIntelligenceRequest) -> dict:
    required_keys = {
        "candidate_profile_type",
        "career_direction_detected",
        "career_change_detected",
        "recommended_resume_model",
        "resume_length_rule",
        "target_market_strategy",
        "recruiter_positioning",
        "priority_sections",
        "sections_to_minimize_or_remove",
        "skill_grouping_strategy",
        "ats_keyword_strategy",
        "missing_information",
        "content_risk_flags",
        "writing_guidance",
    }

    if not required_keys.issubset(parsed.keys()):
        raise HTTPException(
            status_code=500,
            detail=f"Missing keys in resume intelligence response. Required: {required_keys}. Got: {list(parsed.keys())}",
        )

    strategy = get_resume_intelligence_strategy(data)
    parsed["candidate_profile_type"] = strategy["candidate_profile_type"]
    parsed["career_direction_detected"] = strategy["career_direction_detected"]
    parsed["career_change_detected"] = strategy["career_change_detected"]
    parsed["recommended_resume_model"] = strategy["recommended_resume_model"]
    parsed["resume_length_rule"] = strategy["resume_length_rule"]

    career_change_detected = parsed.get("career_change_detected", False)
    if isinstance(career_change_detected, str):
        parsed["career_change_detected"] = career_change_detected.strip().lower() in {"true", "yes", "1"}
    else:
        parsed["career_change_detected"] = bool(career_change_detected)

    parsed["priority_sections"] = clean_string_list(parsed.get("priority_sections", []))
    parsed["sections_to_minimize_or_remove"] = clean_string_list(parsed.get("sections_to_minimize_or_remove", []))
    parsed["skill_grouping_strategy"] = clean_skill_groups(parsed.get("skill_grouping_strategy", []))
    parsed["ats_keyword_strategy"] = clean_string_list(parsed.get("ats_keyword_strategy", []))
    parsed["missing_information"] = clean_string_list(parsed.get("missing_information", []))
    parsed["content_risk_flags"] = clean_string_list(parsed.get("content_risk_flags", []))
    parsed["target_market_strategy"] = str(parsed.get("target_market_strategy", "")).strip()
    parsed["recruiter_positioning"] = str(parsed.get("recruiter_positioning", "")).strip()
    parsed["writing_guidance"] = str(parsed.get("writing_guidance", "")).strip()

    return parsed


def generate_resume_intelligence(data: ResumeIntelligenceRequest) -> dict:
    country_rules = get_country_rules(data.target_country)
    length_rules = get_resume_length_rules(data.experience_level)
    system_msg = (
        "You are a senior career strategist, ATS specialist, recruiter, and resume architect. "
        "Your task is NOT to write the resume. "
        "Your task is to analyze the candidate profile and create a structured resume strategy. "
        "Think like a professional resume writer before writing begins. "
        "Do not invent facts. "
        "Be specific to target role, country, industry, experience level, and career direction."
    )

    user_msg = f"""
Analyze this candidate for resume strategy.

Target Role: {data.target_role}
Target Country: {data.target_country}
Target Industry: {data.target_industry}
Career Direction: {data.career_direction}
Experience Level: {data.experience_level}

Country Rules Engine:
{country_rules}

Resume Length Rules:
{length_rules}

Resume Style Engine:
{get_resume_style_models()}

Current Background:
{data.current_background}

Highest Qualification:
{data.highest_qualification}

Education Details:
{data.education_details}

Work Experience:
{data.work_experience}

Internships:
{data.internships}

Projects:
{data.projects}

Technical Skills:
{data.technical_skills}

Transferable Skills:
{data.transferable_skills}

Tools / Software:
{data.tools_software}

Certifications:
{data.certifications}

Achievements:
{data.achievements}

Leadership Experience:
{data.leadership_experience}

Career Change:
{data.career_change}

Current Field:
{data.current_field}

Target Field:
{data.target_field}

Return ONLY valid JSON in this exact format:

{{
  "candidate_profile_type": "Student | Fresher | Early Career | Mid-Level | Senior | Executive",
  "career_direction_detected": "Technical | Non-Technical | Management | Sales | Marketing | Operations | Finance | Healthcare | Customer Support | Other",
  "career_change_detected": true,
  "recommended_resume_model": "Graduate ATS Resume | Technical Professional Resume | Career Switcher Resume | Business Professional Resume | Executive Leadership Resume",
  "resume_length_rule": "One Page | One to Two Pages | Two Pages",
  "target_market_strategy": "string",
  "recruiter_positioning": "string",
  "priority_sections": ["section1", "section2"],
  "sections_to_minimize_or_remove": ["section1", "section2"],
  "skill_grouping_strategy": [
    {{
      "category": "string",
      "skills": ["skill1", "skill2"]
    }}
  ],
  "ats_keyword_strategy": ["keyword1", "keyword2"],
  "missing_information": ["missing item 1", "missing item 2"],
  "content_risk_flags": ["risk1", "risk2"],
  "writing_guidance": "string"
}}

Rules:
1. Do not write the resume.
2. Do not create fake achievements.
3. Do not create fake metrics.
4. If experience level is Student, Fresher, or 1-5 years, recommend One Page unless there is strong reason otherwise.
5. If experience is 5-10 years, recommend One Page or One to Two Pages depending on content.
6. If experience is 10+ years, recommend Two Pages.
7. Never output labels like Template: India or Template: UK.
8. Use target country only to shape strategy, not as visible template label.
9. Group skills professionally.
10. Identify missing information that would improve the resume.
11. Identify sections that should be minimized.
12. For career changers, focus on transferable skills.
13. For technical roles, prioritize technical skills, projects, certifications, and relevant experience.
14. For non-technical roles, prioritize communication, operations, business tools, customer handling, leadership, coordination, reporting, and transferable skills.
15. Keep recruiter positioning specific and market-aware without writing the final resume.
"""

    resp = client.chat.completions.create(
        model="gpt-4.1-mini",
        messages=[
            {"role": "system", "content": system_msg},
            {"role": "user", "content": user_msg},
        ],
        temperature=0.1,
    )

    content = (resp.choices[0].message.content or "").strip()
    parsed = parse_json_response(content)
    return normalize_resume_intelligence_result(parsed, data)


def build_resume_intelligence_request(data: ResumeIntelligenceInput) -> ResumeIntelligenceRequest:
    return ResumeIntelligenceRequest(
        target_role=data.target_role,
        target_country=data.target_country,
        target_industry=data.target_industry,
        career_direction=data.career_direction,
        experience_level=data.experience_level,
        current_background=data.current_background,
        highest_qualification=data.highest_qualification,
        education_details=data.education_details,
        work_experience=data.work_experience,
        internships=data.internships,
        projects=data.projects,
        technical_skills=data.technical_skills,
        transferable_skills=data.transferable_skills,
        tools_software=data.tools_software,
        certifications=data.certifications,
        achievements=data.achievements,
        leadership_experience=data.leadership_experience,
        career_change=data.career_change,
        current_field=data.current_field,
        target_field=data.target_field,
    )


def get_resume_section_order(resume_style: str) -> str:
    normalized = (resume_style or "").strip().lower()
    if "graduate" in normalized:
        return "Summary -> Education -> Projects -> Skills -> Internships -> Certifications -> Experience"
    if "technical" in normalized:
        return "Summary -> Core Skills -> Professional Experience -> Projects -> Certifications -> Education"
    if "business" in normalized or "executive" in normalized or "career switcher" in normalized:
        return "Summary -> Core Competencies -> Professional Experience -> Achievements -> Education"
    return "Summary -> Skills -> Experience -> Projects -> Education -> Certifications"


def review_resume_quality(resume_text, candidate_data, intelligence, skill_intelligence):
    text = str(resume_text or "")
    lowered = text.lower()
    issues_found = []
    required_fixes = []

    generic_phrases = [
        "motivated and detail-oriented",
        "hardworking individual",
        "seeking an opportunity",
        "passionate about learning",
        "team player with good communication skills",
        "eager to contribute",
    ]
    placeholder_markers = [
        "[no formal work experience]",
        "[details not provided]",
        "[no certifications provided]",
        "[year not provided]",
        "[current employer]",
        "[dates]",
        "[institution]",
        "[company]",
        "not provided",
    ]

    quality_score = 100
    ats_score = 100
    readability_score = 100
    role_alignment_score = 100

    summary_text = ""
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    for idx, line in enumerate(lines):
        if line.lower() in {"professional summary", "executive summary", "summary"} and idx + 1 < len(lines):
            collected = []
            for next_line in lines[idx + 1:]:
                if next_line.isupper() or next_line.lower() in {"education", "projects", "skills", "technical skills", "core skills", "core competencies", "experience", "professional experience", "internships", "certifications", "achievements"}:
                    break
                collected.append(next_line)
            summary_text = " ".join(collected).strip()
            break

    if not summary_text:
        summary_text = text[:500]

    summary_lower = summary_text.lower()
    if any(phrase in summary_lower for phrase in generic_phrases):
        issues_found.append("Professional summary uses generic AI-style phrasing.")
        required_fixes.append("Rewrite the summary with sharper role positioning and natural language.")
        quality_score -= 18
        readability_score -= 10

    if len(summary_text.split()) < 45 or len(summary_text.split()) > 140:
        issues_found.append("Professional summary length is outside the preferred premium range.")
        required_fixes.append("Keep the summary concise, role-specific, and around 70-120 words.")
        quality_score -= 8
        readability_score -= 5

    placeholder_lines = [line.strip().lower() for line in text.splitlines() if line.strip()]
    has_placeholder = any(marker in lowered for marker in placeholder_markers) or any(line in {"n/a", "na", "not provided"} for line in placeholder_lines)
    if has_placeholder:
        issues_found.append("Resume contains placeholders or missing-information markers.")
        required_fixes.append("Remove placeholders and hide incomplete sections instead.")
        quality_score -= 30
        ats_score -= 25
        readability_score -= 20

    role_words = [word for word in str(getattr(candidate_data, "target_role", "")).lower().replace("/", " ").split() if len(word) > 2]
    role_hits = sum(1 for word in role_words if word in lowered)
    if role_words and role_hits == 0:
        issues_found.append("Resume does not clearly align to the target role.")
        required_fixes.append("Make the title, summary, and core content more role-specific.")
        quality_score -= 18
        role_alignment_score -= 22
    elif role_words and role_hits < max(1, len(role_words) // 2):
        role_alignment_score -= 10
        quality_score -= 6

    expected_keywords = clean_string_list(intelligence.get("ats_keyword_strategy", []))
    keyword_hits = sum(1 for keyword in expected_keywords if keyword.lower() in lowered)
    if expected_keywords:
        keyword_ratio = keyword_hits / max(1, len(expected_keywords))
        if keyword_ratio < 0.35:
            issues_found.append("Target-role keywords are underrepresented.")
            required_fixes.append("Integrate more role-relevant keywords naturally across summary, skills, and experience.")
            ats_score -= 18
            role_alignment_score -= 12
            quality_score -= 10
        elif keyword_ratio < 0.55:
            ats_score -= 8
            quality_score -= 4

    expected_groups = clean_skill_groups(skill_intelligence.get("skill_groups", []))
    grouped_categories_present = sum(1 for group in expected_groups if group["category"].lower() in lowered)
    if expected_groups and grouped_categories_present < max(2, len(expected_groups) // 2):
        issues_found.append("Skills are not clearly grouped using the Skill Intelligence structure.")
        required_fixes.append("Rebuild the skills section using the provided grouped categories instead of a flat list.")
        quality_score -= 15
        ats_score -= 10

    repeated_phrases = []
    for phrase in ["proven ability", "strong ability", "responsible for", "team coordination", "service quality"]:
        if lowered.count(phrase) > 2:
            repeated_phrases.append(phrase)
    if repeated_phrases:
        issues_found.append("Resume contains repeated phrasing that weakens readability.")
        required_fixes.append("Vary sentence openings and remove repetitive phrasing.")
        quality_score -= 8
        readability_score -= 10

    expected_order = get_resume_section_order(intelligence.get("recommended_resume_model", ""))
    expected_sections = [section.strip().lower() for section in expected_order.split("->")]
    aliases = {
        "summary": ["professional summary", "executive summary", "summary"],
        "education": ["education"],
        "projects": ["projects", "project experience"],
        "skills": ["skills", "technical skills", "core skills", "core competencies", "skills & certifications"],
        "internships": ["internships", "internship experience"],
        "certifications": ["certifications", "licenses & certifications"],
        "experience": ["experience", "professional experience", "relevant experience"],
        "core skills": ["core skills", "technical skills"],
        "core competencies": ["core competencies", "skills", "skills & certifications"],
        "professional experience": ["professional experience", "experience", "relevant experience"],
        "achievements": ["achievements", "key achievements", "leadership & achievements"],
    }
    positions = []
    for section in expected_sections:
        section_positions = []
        for alias in aliases.get(section, [section]):
            pos = lowered.find(alias)
            if pos != -1:
                section_positions.append(pos)
        if section_positions:
            positions.append(min(section_positions))
    if positions and positions != sorted(positions):
        issues_found.append("Section order does not fully follow the recommended resume model.")
        required_fixes.append("Reorder the resume sections to match the recommended structure.")
        quality_score -= 7
        readability_score -= 6

    word_count = len(text.split())
    length_rule = str(intelligence.get("resume_length_rule", ""))
    if length_rule == "One Page" and word_count > 900:
        issues_found.append("Resume is too long for the recommended one-page format.")
        required_fixes.append("Tighten the resume and remove lower-priority detail to fit a one-page style.")
        quality_score -= 10
        readability_score -= 8
    elif length_rule == "Two Pages" and word_count < 280:
        issues_found.append("Resume may be too thin for the recommended senior-level two-page style.")
        required_fixes.append("Strengthen leadership, scope, and impact detail where supported by the input.")
        quality_score -= 6
        role_alignment_score -= 4

    recruiter_value_fast = bool(str(getattr(candidate_data, "target_role", "")).strip()) and bool(summary_text.strip()) and len(summary_text.split()) >= 45
    if not recruiter_value_fast:
        issues_found.append("Candidate value is not immediately clear to a recruiter.")
        required_fixes.append("Sharpen the professional title and summary so the target value is clear within seconds.")
        quality_score -= 12
        readability_score -= 8

    quality_score = max(0, min(100, quality_score))
    ats_score = max(0, min(100, ats_score))
    readability_score = max(0, min(100, readability_score))
    role_alignment_score = max(0, min(100, role_alignment_score))

    if quality_score >= 88:
        summary_quality = "Strong"
    elif quality_score >= 72:
        summary_quality = "Average"
    else:
        summary_quality = "Weak"

    is_ready_for_user = quality_score >= 80 and not has_placeholder

    return {
        "quality_score": quality_score,
        "ats_readiness_score": ats_score,
        "recruiter_readability_score": readability_score,
        "role_alignment_score": role_alignment_score,
        "summary_quality": summary_quality,
        "issues_found": clean_string_list(issues_found),
        "required_fixes": clean_string_list(required_fixes),
        "is_ready_for_user": is_ready_for_user,
    }


def normalize_build_resume_response(parsed: dict, intelligence: dict, skill_intelligence: dict, quality_report: dict, quality_fixes_applied: list[str]) -> dict:
    parsed["recommended_resume_style"] = str(parsed.get("recommended_resume_style", "")).strip() or intelligence["recommended_resume_model"]
    parsed["recommendation_reason"] = str(parsed.get("recommendation_reason", "")).strip()
    parsed["professional_title"] = str(parsed.get("professional_title", "")).strip()
    parsed["executive_summary"] = str(parsed.get("executive_summary", "")).strip()
    parsed["resume_length_rule"] = str(parsed.get("resume_length_rule") or intelligence["resume_length_rule"]).strip() or intelligence["resume_length_rule"]
    parsed["target_market_strategy"] = str(parsed.get("target_market_strategy") or intelligence["target_market_strategy"]).strip() or intelligence["target_market_strategy"]
    parsed["recruiter_positioning"] = str(parsed.get("recruiter_positioning") or intelligence["recruiter_positioning"]).strip() or intelligence["recruiter_positioning"]
    parsed["full_resume"] = sanitize_resume_text(str(parsed.get("full_resume", "")).strip())
    parsed["ats_keywords"] = clean_string_list(parsed.get("ats_keywords", [])) or intelligence["ats_keyword_strategy"]
    parsed["skill_groups"] = clean_skill_groups(skill_intelligence["skill_groups"])
    parsed["strengths"] = clean_string_list(parsed.get("strengths", []))
    parsed["missing_information"] = clean_string_list(parsed.get("missing_information", [])) or intelligence["missing_information"]
    parsed["improvement_suggestions"] = clean_string_list(parsed.get("improvement_suggestions", []))
    parsed["writing_quality_score"] = str(parsed.get("writing_quality_score", "")).strip()
    parsed["resume_readability"] = str(parsed.get("resume_readability", "")).strip()
    parsed["ats_readiness"] = str(parsed.get("ats_readiness", "")).strip()
    parsed["resume_confidence"] = str(parsed.get("resume_confidence", "")).strip()
    parsed["quality_score"] = int(quality_report["quality_score"])
    parsed["ats_readiness_score"] = int(quality_report["ats_readiness_score"])
    parsed["recruiter_readability_score"] = int(quality_report["recruiter_readability_score"])
    parsed["role_alignment_score"] = int(quality_report["role_alignment_score"])
    parsed["quality_issues_found"] = clean_string_list(quality_report.get("issues_found", []))
    parsed["quality_fixes_applied"] = clean_string_list(quality_fixes_applied)

    if not parsed["writing_quality_score"]:
        parsed["writing_quality_score"] = f"{quality_report['quality_score']}/100"
    if not parsed["resume_readability"]:
        readability = int(quality_report["recruiter_readability_score"])
        parsed["resume_readability"] = "High" if readability >= 85 else "Moderate" if readability >= 70 else "Needs improvement"
    if not parsed["ats_readiness"]:
        ats = int(quality_report["ats_readiness_score"])
        parsed["ats_readiness"] = "High" if ats >= 85 else "Moderate" if ats >= 70 else "Needs improvement"
    if not parsed["resume_confidence"]:
        confidence = int(quality_report["quality_score"])
        parsed["resume_confidence"] = "Strong" if confidence >= 85 else "Promising" if confidence >= 70 else "Developing"

    if skill_intelligence["missing_role_skills"]:
        parsed["improvement_suggestions"].append(
            f"Consider building stronger exposure in: {', '.join(skill_intelligence['missing_role_skills'])}."
        )
        parsed["improvement_suggestions"] = clean_string_list(parsed["improvement_suggestions"])

    return parsed


def normalize_recruiter_review(parsed: dict) -> dict:
    required_keys = {
        "interview_probability",
        "recruiter_confidence",
        "first_impression",
        "shortlisting_decision",
        "top_strengths",
        "top_concerns",
        "missing_high_value_information",
        "recommended_improvements",
        "industry_keywords_missing",
        "resume_competitiveness",
    }
    if not required_keys.issubset(parsed.keys()):
        raise HTTPException(
            status_code=500,
            detail=f"Missing keys in recruiter review response. Required: {required_keys}. Got: {list(parsed.keys())}",
        )

    def clamp_score(value):
        try:
            return max(0, min(100, int(value)))
        except Exception:
            return 0

    parsed["interview_probability"] = clamp_score(parsed.get("interview_probability", 0))
    parsed["recruiter_confidence"] = clamp_score(parsed.get("recruiter_confidence", 0))
    parsed["first_impression"] = str(parsed.get("first_impression", "")).strip()
    decision = str(parsed.get("shortlisting_decision", "Maybe")).strip().title()
    if decision not in {"Yes", "Maybe", "No"}:
        decision = "Maybe"
    parsed["shortlisting_decision"] = decision
    parsed["top_strengths"] = clean_string_list(parsed.get("top_strengths", []))
    parsed["top_concerns"] = clean_string_list(parsed.get("top_concerns", []))
    parsed["missing_high_value_information"] = clean_string_list(parsed.get("missing_high_value_information", []))
    parsed["recommended_improvements"] = clean_string_list(parsed.get("recommended_improvements", []))
    parsed["industry_keywords_missing"] = clean_string_list(parsed.get("industry_keywords_missing", []))
    competitiveness = str(parsed.get("resume_competitiveness", "Competitive")).strip().title()
    if competitiveness not in {"Basic", "Competitive", "Strong", "Outstanding"}:
        competitiveness = "Competitive"
    parsed["resume_competitiveness"] = competitiveness
    return parsed


def normalize_job_description_analysis(parsed: dict) -> dict:
    required_keys = {
        "job_level",
        "career_direction",
        "industry",
        "experience_required",
        "education_expectation",
        "required_skills",
        "preferred_skills",
        "technical_skills",
        "soft_skills",
        "tools_and_platforms",
        "ats_keywords",
        "primary_responsibilities",
        "leadership_required",
        "communication_level",
        "problem_solving_level",
        "resume_focus_sections",
        "recommended_resume_model",
        "important_certifications",
        "company_type_guess",
        "hiring_priority_summary",
        "candidate_fit_strategy",
    }
    if not required_keys.issubset(parsed.keys()):
        raise HTTPException(
            status_code=500,
            detail=f"Missing keys in job description analysis response. Required: {required_keys}. Got: {list(parsed.keys())}",
        )

    parsed["required_skills"] = clean_string_list(parsed.get("required_skills", []))
    parsed["preferred_skills"] = clean_string_list(parsed.get("preferred_skills", []))
    parsed["technical_skills"] = clean_string_list(parsed.get("technical_skills", []))
    parsed["soft_skills"] = clean_string_list(parsed.get("soft_skills", []))
    parsed["tools_and_platforms"] = clean_string_list(parsed.get("tools_and_platforms", []))
    parsed["ats_keywords"] = clean_string_list(parsed.get("ats_keywords", []))
    parsed["primary_responsibilities"] = clean_string_list(parsed.get("primary_responsibilities", []))
    parsed["resume_focus_sections"] = clean_string_list(parsed.get("resume_focus_sections", []))
    parsed["important_certifications"] = clean_string_list(parsed.get("important_certifications", []))
    parsed["leadership_required"] = bool(parsed.get("leadership_required", False))
    parsed["job_level"] = str(parsed.get("job_level", "")).strip()
    parsed["career_direction"] = str(parsed.get("career_direction", "")).strip()
    parsed["industry"] = str(parsed.get("industry", "")).strip()
    parsed["experience_required"] = str(parsed.get("experience_required", "")).strip()
    parsed["education_expectation"] = str(parsed.get("education_expectation", "")).strip()
    parsed["communication_level"] = str(parsed.get("communication_level", "")).strip()
    parsed["problem_solving_level"] = str(parsed.get("problem_solving_level", "")).strip()
    parsed["recommended_resume_model"] = str(parsed.get("recommended_resume_model", "")).strip()
    parsed["company_type_guess"] = str(parsed.get("company_type_guess", "")).strip()
    parsed["hiring_priority_summary"] = str(parsed.get("hiring_priority_summary", "")).strip()
    parsed["candidate_fit_strategy"] = str(parsed.get("candidate_fit_strategy", "")).strip()
    return parsed


def analyze_job_description_intelligence(data: JobDescriptionRequest) -> dict:
    career_knowledge = generate_career_knowledge(
        {
            "target_role": data.job_title,
            "target_industry": data.industry,
            "education": "",
            "highest_qualification": "",
            "current_background": data.job_description,
            "career_direction": "",
            "experience_level": "",
        }
    )
    system_msg = (
        "You are a Senior Recruiter, Hiring Manager, ATS Specialist, and Career Strategist. "
        "Analyze the job description and extract structured intelligence about what the employer wants. "
        "Do not rewrite the job description. Do not return generic summaries. Return only structured intelligence."
    )

    user_msg = f"""
Analyze this job description and return structured job intelligence.

Job title: {data.job_title}
Company name: {data.company_name}
Country: {data.country}
Industry: {data.industry}
Career knowledge graph:
Recommended roles: {career_knowledge['recommended_roles']}
Recommended industries: {career_knowledge['recommended_industries']}
Recommended certifications: {career_knowledge['recommended_certifications']}
Job description:
{data.job_description}

Return ONLY valid JSON in this exact format:
{{
"job_level": "",
"career_direction": "",
"industry": "",
"experience_required": "",
"education_expectation": "",
"required_skills": [],
"preferred_skills": [],
"technical_skills": [],
"soft_skills": [],
"tools_and_platforms": [],
"ats_keywords": [],
"primary_responsibilities": [],
"leadership_required": false,
"communication_level": "",
"problem_solving_level": "",
"resume_focus_sections": [],
"recommended_resume_model": "",
"important_certifications": [],
"company_type_guess": "",
"hiring_priority_summary": "",
"candidate_fit_strategy": ""
}}
"""

    resp = client.chat.completions.create(
        model="gpt-4.1-mini",
        messages=[
            {"role": "system", "content": system_msg},
            {"role": "user", "content": user_msg},
        ],
        temperature=0.1,
    )
    content = (resp.choices[0].message.content or "").strip()
    parsed = parse_json_response(content)
    return normalize_job_description_analysis(parsed)


def compare_candidate_to_job_description(candidate_data, job_intelligence: dict) -> dict:
    candidate_skills = split_skills_text(
        getattr(candidate_data, "technical_skills", ""),
        getattr(candidate_data, "transferable_skills", ""),
        getattr(candidate_data, "tools_software", ""),
        getattr(candidate_data, "certifications", ""),
        getattr(candidate_data, "projects", ""),
        getattr(candidate_data, "work_experience", ""),
        getattr(candidate_data, "internships", ""),
    )
    candidate_text = "\n".join(
        [
            getattr(candidate_data, "technical_skills", ""),
            getattr(candidate_data, "transferable_skills", ""),
            getattr(candidate_data, "tools_software", ""),
            getattr(candidate_data, "certifications", ""),
            getattr(candidate_data, "projects", ""),
            getattr(candidate_data, "work_experience", ""),
            getattr(candidate_data, "internships", ""),
            getattr(candidate_data, "resume_text", ""),
        ]
    ).lower()
    candidate_skill_map = {skill.lower(): skill for skill in candidate_skills}

    def match_against(target_skills: list[str]) -> tuple[list[str], list[str]]:
        strong = []
        missing = []
        for target in clean_string_list(target_skills):
            target_lower = target.lower()
            matched = None
            for candidate_lower, original in candidate_skill_map.items():
                if target_lower in candidate_lower or candidate_lower in target_lower:
                    matched = original
                    break
            if not matched and target_lower in candidate_text:
                matched = target
            if matched:
                strong.append(title_case_skill(matched))
            else:
                missing.append(target)
        return clean_string_list(strong), clean_string_list(missing)

    strong_required, missing_required = match_against(job_intelligence.get("required_skills", []))
    strong_preferred, missing_preferred = match_against(job_intelligence.get("preferred_skills", []))
    strong_matching_skills = clean_string_list(strong_required + strong_preferred)

    total_targets = len(clean_string_list(job_intelligence.get("required_skills", []))) + len(clean_string_list(job_intelligence.get("preferred_skills", [])))
    total_matches = len(set(skill.lower() for skill in strong_matching_skills))
    if total_targets == 0:
        skill_match_percentage = 0
    else:
        skill_match_percentage = int(round((total_matches / total_targets) * 100))

    if missing_required:
        resume_alignment_strategy = "Lead with the strongest matching skills, keep role-relevant evidence visible early, and address the missing required skills through truthful positioning or future upskilling suggestions."
    elif missing_preferred:
        resume_alignment_strategy = "Emphasize the strong core fit and align the resume language closely to employer priorities while treating preferred gaps as secondary development areas."
    else:
        resume_alignment_strategy = "The candidate already aligns well with the stated requirements, so the resume should focus on clarity, proof of capability, and ATS keyword coverage."

    return {
        "skill_match_percentage": max(0, min(100, skill_match_percentage)),
        "strong_matching_skills": strong_matching_skills[:12],
        "missing_required_skills": missing_required[:12],
        "missing_preferred_skills": missing_preferred[:12],
        "resume_alignment_strategy": resume_alignment_strategy,
    }


def normalize_achievement_intelligence(parsed: dict) -> dict:
    required_keys = {
        "experience_bullets",
        "project_bullets",
        "leadership_bullets",
        "transferable_achievement_bullets",
        "missing_impact_questions",
    }
    if not required_keys.issubset(parsed.keys()):
        raise HTTPException(
            status_code=500,
            detail=f"Missing keys in achievement intelligence response. Required: {required_keys}. Got: {list(parsed.keys())}",
        )

    cleaned_experience = []
    for item in parsed.get("experience_bullets", []):
        if not isinstance(item, dict):
            continue
        original = str(item.get("original", "")).strip()
        improved = str(item.get("improved", "")).strip()
        impact_type = str(item.get("impact_type", "")).strip()
        confidence = str(item.get("confidence", "")).strip()
        if original or improved:
            cleaned_experience.append(
                {
                    "original": original,
                    "improved": improved,
                    "impact_type": impact_type,
                    "confidence": confidence,
                }
            )

    cleaned_projects = []
    for item in parsed.get("project_bullets", []):
        if not isinstance(item, dict):
            continue
        project_name = str(item.get("project_name", "")).strip()
        improved_bullets = clean_string_list(item.get("improved_bullets", []))
        technologies_detected = clean_string_list(item.get("technologies_detected", []))
        project_value = str(item.get("project_value", "")).strip()
        if project_name or improved_bullets:
            cleaned_projects.append(
                {
                    "project_name": project_name,
                    "improved_bullets": improved_bullets,
                    "technologies_detected": technologies_detected,
                    "project_value": project_value,
                }
            )

    parsed["experience_bullets"] = cleaned_experience
    parsed["project_bullets"] = cleaned_projects
    parsed["leadership_bullets"] = clean_string_list(parsed.get("leadership_bullets", []))
    parsed["transferable_achievement_bullets"] = clean_string_list(parsed.get("transferable_achievement_bullets", []))
    parsed["missing_impact_questions"] = clean_string_list(parsed.get("missing_impact_questions", []))
    return parsed


def generate_achievement_intelligence(candidate_data, intelligence=None, job_intelligence=None):
    intelligence = intelligence or {}
    job_intelligence = job_intelligence or {}
    system_msg = (
        "You are an Achievement Intelligence Engine for professional resumes. "
        "Convert weak responsibilities, projects, internships, and leadership statements into strong, truthful, achievement-oriented bullets. "
        "Do not invent fake numbers, fake metrics, fake companies, fake dates, fake tools, or fake achievements. "
        "If impact is implied but not measured, use honest impact-based wording without exaggeration."
    )

    user_msg = f"""
Generate achievement intelligence from this candidate information.

Target role: {getattr(candidate_data, 'target_role', '')}
Target country: {getattr(candidate_data, 'target_country', 'Global')}
Experience level: {getattr(candidate_data, 'experience_level', '')}
Career direction: {getattr(candidate_data, 'career_direction', '')}

Resume intelligence:
Recommended resume model: {intelligence.get('recommended_resume_model', '')}
Recruiter positioning: {intelligence.get('recruiter_positioning', '')}
Priority sections: {json.dumps(intelligence.get('priority_sections', []))}

Job description intelligence:
{json.dumps(job_intelligence) if job_intelligence else 'Not provided'}

Raw experience content:
Work experience:
{getattr(candidate_data, 'work_experience', getattr(candidate_data, 'resume_text', ''))}

Internships:
{getattr(candidate_data, 'internships', '')}

Projects:
{getattr(candidate_data, 'projects', getattr(candidate_data, 'resume_text', ''))}

Achievements:
{getattr(candidate_data, 'achievements', '')}

Leadership experience:
{getattr(candidate_data, 'leadership_experience', '')}

Technical skills:
{getattr(candidate_data, 'technical_skills', '')}

Transferable skills:
{getattr(candidate_data, 'transferable_skills', '')}

Return ONLY valid JSON in this exact format:
{{
  "experience_bullets": [
    {{
      "original": "string",
      "improved": "string",
      "impact_type": "Technical | Business | Operational | Customer | Leadership | Academic | Transferable",
      "confidence": "High | Medium | Low"
    }}
  ],
  "project_bullets": [
    {{
      "project_name": "string",
      "improved_bullets": ["bullet1", "bullet2"],
      "technologies_detected": ["tech1", "tech2"],
      "project_value": "string"
    }}
  ],
  "leadership_bullets": ["bullet1", "bullet2"],
  "transferable_achievement_bullets": ["bullet1", "bullet2"],
  "missing_impact_questions": ["question1", "question2"]
}}

Rules:
1. Do not invent numbers.
2. If no measurable result is provided, use impact-based wording without fake metrics.
3. Convert responsibilities into achievement-style bullets.
4. Use strong action verbs.
5. Avoid generic wording.
6. Avoid phrases like responsible for, worked on, helped with, involved in.
7. Do not overstate fresher experience.
8. For students and freshers, convert projects, internships, and academic work into professional bullets.
9. For career switchers, convert previous experience into transferable achievements.
10. For senior professionals, emphasize leadership, ownership, strategic impact, team management, operations, delivery, and process improvement where supported.
"""

    resp = client.chat.completions.create(
        model="gpt-4.1-mini",
        messages=[
            {"role": "system", "content": system_msg},
            {"role": "user", "content": user_msg},
        ],
        temperature=0.15,
    )
    content = (resp.choices[0].message.content or "").strip()
    parsed = parse_json_response(content)
    return normalize_achievement_intelligence(parsed)


def generate_ats_intelligence(candidate_data, resume_text="", job_intelligence=None, intelligence=None, skill_intelligence=None):
    job_intelligence = job_intelligence or {}
    intelligence = intelligence or {}
    skill_intelligence = skill_intelligence or {}

    target_role = str(getattr(candidate_data, "target_role", "")).strip()
    target_country = str(getattr(candidate_data, "target_country", "Global")).strip()
    target_industry = str(getattr(candidate_data, "target_industry", "")).strip()
    experience_level = str(getattr(candidate_data, "experience_level", "")).strip()
    career_direction = str(getattr(candidate_data, "career_direction", "")).strip()

    role_targets = get_role_skill_targets(target_role, career_direction, target_industry)
    library_keywords = []
    for _, keywords in role_targets.get("categories", []):
        library_keywords.extend(keywords)
    library_keywords.extend(role_targets.get("missing", []))

    if job_intelligence:
        required_keywords = clean_string_list(
            job_intelligence.get("required_skills", [])
            + job_intelligence.get("technical_skills", [])
            + job_intelligence.get("tools_and_platforms", [])
            + job_intelligence.get("soft_skills", [])
            + job_intelligence.get("ats_keywords", [])
        )
    else:
        required_keywords = clean_string_list(
            intelligence.get("ats_keyword_strategy", [])
            + library_keywords
            + [target_role, target_industry]
        )

    deduped_required = []
    seen_required = set()
    for keyword in required_keywords:
        key = keyword.lower()
        if not key or key in seen_required:
            continue
        seen_required.add(key)
        deduped_required.append(title_case_skill(keyword))
    required_keywords = deduped_required[:18]

    candidate_text = "\n".join(
        [
            resume_text,
            getattr(candidate_data, "technical_skills", ""),
            getattr(candidate_data, "transferable_skills", ""),
            getattr(candidate_data, "tools_software", ""),
            getattr(candidate_data, "projects", ""),
            getattr(candidate_data, "work_experience", ""),
        ]
    ).lower()
    candidate_skill_pool = split_skills_text(
        getattr(candidate_data, "technical_skills", ""),
        getattr(candidate_data, "transferable_skills", ""),
        getattr(candidate_data, "tools_software", ""),
    )
    candidate_skill_map = {skill.lower(): title_case_skill(skill) for skill in candidate_skill_pool}

    matching_keywords = []
    missing_keywords = []
    for keyword in required_keywords:
        keyword_lower = keyword.lower()
        matched_value = None
        for candidate_lower, original in candidate_skill_map.items():
            if keyword_lower in candidate_lower or candidate_lower in keyword_lower:
                matched_value = original
                break
        if not matched_value and keyword_lower and keyword_lower in candidate_text:
            matched_value = keyword
        if matched_value:
            matching_keywords.append(title_case_skill(keyword if keyword_lower in candidate_text else matched_value))
        else:
            missing_keywords.append(keyword)

    matching_keywords = clean_string_list(matching_keywords)
    missing_keywords = clean_string_list(missing_keywords)

    def dedupe_preserve(values: list[str]) -> list[str]:
        cleaned = []
        seen = set()
        for value in clean_string_list(values):
            key = value.lower()
            if key in seen:
                continue
            seen.add(key)
            cleaned.append(value)
        return cleaned

    if not resume_text:
        base_summary = matching_keywords[:4]
        base_skills = matching_keywords[:6]
        base_experience = matching_keywords[:4]
        base_projects = matching_keywords[:4]
    else:
        lowered_resume = resume_text.lower()
        present = [keyword for keyword in required_keywords if keyword.lower() in lowered_resume]
        absent_but_owned = [keyword for keyword in matching_keywords if keyword.lower() not in lowered_resume]
        base_summary = clean_string_list(present[:3] + absent_but_owned[:2])
        base_skills = clean_string_list(present[:4] + absent_but_owned[:3])
        base_experience = clean_string_list(present[:4])
        base_projects = clean_string_list(
            [keyword for keyword in present if any(token in keyword.lower() for token in ["python", "sql", "docker", "linux", "verilog", "cadence", "aws", "excel", "power bi", "analysis", "design"])]
        )[:4]

    if experience_level.strip().lower() in {"student", "fresher", "1-3 years", "1â€“3 years"}:
        keyword_placement_strategy = {
            "summary": dedupe_preserve(base_summary)[:4],
            "skills": dedupe_preserve(base_skills)[:6],
            "experience": dedupe_preserve(base_experience)[:3],
            "projects": dedupe_preserve(base_projects[:4] + matching_keywords[:2])[:5],
        }
    else:
        keyword_placement_strategy = {
            "summary": dedupe_preserve(base_summary)[:3],
            "skills": dedupe_preserve(base_skills)[:6],
            "experience": dedupe_preserve(base_experience[:5] + matching_keywords[:2])[:6],
            "projects": dedupe_preserve(base_projects)[:4],
        }

    formatting_risks = []
    if resume_text:
        for marker, message in [
            ("|", "Possible table-style formatting detected; ATS parsers can misread table layouts."),
            ("•", "Decorative bullet characters may reduce ATS consistency in some systems."),
            ("★", "Decorative symbols or icons can create ATS parsing risk."),
            ("\t", "Tab-based alignment can create ATS formatting issues."),
        ]:
            if marker in resume_text:
                formatting_risks.append(message)
        if any(token in resume_text.lower() for token in ["header", "footer", "text box", "icon"]):
            formatting_risks.append("Resume may rely on headers, footers, icons, or text-box style content that ATS tools can miss.")
    formatting_risks = clean_string_list(formatting_risks)

    lowered_resume = str(resume_text or "").lower()
    section_risks = []
    standard_headings = ["professional summary", "skills", "experience", "projects", "education", "certifications"]
    if resume_text:
        if not any(heading in lowered_resume for heading in ["professional summary", "executive summary", "summary"]):
            section_risks.append("Professional Summary heading is missing or non-standard.")
        if not any(heading in lowered_resume for heading in ["skills", "core skills", "core competencies", "technical skills"]):
            section_risks.append("Skills section is missing or not clearly labeled.")
        if not any(heading in lowered_resume for heading in ["experience", "professional experience", "relevant experience"]):
            section_risks.append("Experience section is missing or not clearly labeled.")
        if getattr(candidate_data, "projects", "") and "projects" not in lowered_resume:
            section_risks.append("Projects content exists but the Projects section is not clearly surfaced.")
        if getattr(candidate_data, "technical_skills", "") and skill_intelligence and len(clean_skill_groups(skill_intelligence.get("skill_groups", []))) == 0:
            section_risks.append("Skills are not grouped clearly, which can weaken ATS and recruiter readability.")
    else:
        if experience_level.strip().lower() in {"student", "fresher"}:
            section_risks.append("For fresher resumes, Projects and Skills sections should appear prominently for ATS alignment.")
    section_risks = clean_string_list(section_risks)

    keyword_stuffing_risk = "Low"
    if resume_text:
        overused = 0
        for keyword in required_keywords[:12]:
            count = lowered_resume.count(keyword.lower())
            if count >= 5:
                overused += 1
        if overused >= 3:
            keyword_stuffing_risk = "High"
        elif overused >= 1:
            keyword_stuffing_risk = "Medium"

    total_required = len(required_keywords)
    match_ratio = len(set(keyword.lower() for keyword in matching_keywords)) / max(1, total_required)
    base_score = int(round(match_ratio * 100))
    ats_score_estimate = base_score
    ats_score_estimate -= min(12, len(formatting_risks) * 4)
    ats_score_estimate -= min(12, len(section_risks) * 4)
    if keyword_stuffing_risk == "Medium":
        ats_score_estimate -= 6
    elif keyword_stuffing_risk == "High":
        ats_score_estimate -= 14
    ats_score_estimate = max(35 if matching_keywords else 20, min(98, ats_score_estimate))

    if ats_score_estimate >= 90:
        ats_readiness_level = "Excellent"
    elif ats_score_estimate >= 78:
        ats_readiness_level = "High"
    elif ats_score_estimate >= 60:
        ats_readiness_level = "Medium"
    else:
        ats_readiness_level = "Low"

    ats_improvement_actions = []
    if missing_keywords:
        ats_improvement_actions.append(
            f"Do not add unsupported skills; instead strengthen truthful evidence around existing strengths and treat these as missing: {', '.join(missing_keywords[:6])}."
        )
    if keyword_placement_strategy["summary"]:
        ats_improvement_actions.append(
            f"Place the strongest role keywords naturally in the Professional Summary: {', '.join(keyword_placement_strategy['summary'][:4])}."
        )
    if keyword_placement_strategy["skills"]:
        ats_improvement_actions.append(
            f"Keep the Skills section ATS-safe and grouped, highlighting: {', '.join(keyword_placement_strategy['skills'][:5])}."
        )
    if formatting_risks:
        ats_improvement_actions.append("Use simple ATS-safe formatting with standard headings and avoid tables, icons, columns, and decorative layouts.")
    if section_risks:
        ats_improvement_actions.append("Use standard section titles such as Professional Summary, Skills, Experience, Projects, Education, and Certifications.")
    if job_intelligence:
        ats_strategy_note = "Prioritize employer keywords from the job description, but only where the candidate has real supporting evidence."
    else:
        ats_strategy_note = "Use target-role and industry terminology naturally across summary, grouped skills, projects, and experience without keyword stuffing."
    if experience_level.strip().lower() in {"student", "fresher"}:
        ats_strategy_note += " For fresher profiles, concentrate more ATS evidence in Summary, Skills, Projects, and Education."
    elif str(getattr(candidate_data, "career_change", "No")).strip().lower() in {"yes", "true", "1"}:
        ats_strategy_note += " For career-change profiles, emphasize transferable language that is supported by real experience."

    return {
        "ats_score_estimate": ats_score_estimate,
        "ats_readiness_level": ats_readiness_level,
        "required_keywords": required_keywords,
        "matching_keywords": dedupe_preserve(matching_keywords)[:12],
        "missing_keywords": dedupe_preserve(missing_keywords)[:12],
        "keyword_placement_strategy": {
            "summary": clean_string_list(keyword_placement_strategy["summary"])[:6],
            "skills": clean_string_list(keyword_placement_strategy["skills"])[:8],
            "experience": clean_string_list(keyword_placement_strategy["experience"])[:6],
            "projects": clean_string_list(keyword_placement_strategy["projects"])[:6],
        },
        "formatting_risks": formatting_risks[:8],
        "section_risks": section_risks[:8],
        "keyword_stuffing_risk": keyword_stuffing_risk,
        "ats_improvement_actions": clean_string_list(ats_improvement_actions)[:8],
        "ats_strategy_note": ats_strategy_note,
    }


def guess_company_type(candidate_data, job_intelligence=None) -> str:
    job_intelligence = job_intelligence or {}
    guessed = str(job_intelligence.get("company_type_guess", "")).strip()
    if guessed:
        return guessed

    combined = " ".join(
        [
            str(getattr(candidate_data, "target_role", "") or ""),
            str(getattr(candidate_data, "target_industry", "") or ""),
            str(getattr(candidate_data, "job_description", "") or ""),
        ]
    ).lower()

    mapping = [
        ("Government", ["government", "public sector", "compliance", "administration"]),
        ("Research", ["research", "laboratory", "r&d", "innovation", "scientist"]),
        ("Healthcare", ["healthcare", "hospital", "clinical", "medical"]),
        ("Hospitality", ["restaurant", "hospitality", "hotel", "food service"]),
        ("Manufacturing", ["manufacturing", "plant", "production", "factory", "industrial"]),
        ("Retail", ["retail", "store", "merchandising", "customer floor"]),
        ("Consultancy", ["consulting", "consultancy", "advisory", "client delivery"]),
        ("Product Company", ["product", "platform", "saas", "software product"]),
        ("Enterprise", ["enterprise", "global", "large scale", "standards", "cross-functional"]),
        ("Startup", ["startup", "fast-paced", "0-1", "ownership", "wear multiple hats"]),
    ]
    for label, keywords in mapping:
        if any(keyword in combined for keyword in keywords):
            return label
    return "Corporate"


def generate_resume_personalization(candidate_data, job_intelligence=None, intelligence=None, ats_intelligence=None, recruiter_intelligence=None):
    job_intelligence = job_intelligence or {}
    intelligence = intelligence or {}
    ats_intelligence = ats_intelligence or {}
    recruiter_intelligence = recruiter_intelligence or {}

    target_country = str(getattr(candidate_data, "target_country", "Global")).strip().lower()
    target_role = str(getattr(candidate_data, "target_role", "")).strip()
    target_industry = str(getattr(candidate_data, "target_industry", "")).strip()
    experience_level = str(getattr(candidate_data, "experience_level", "")).strip().lower()
    role_direction = str(
        getattr(candidate_data, "career_direction", "")
        or intelligence.get("career_direction_detected", "")
        or job_intelligence.get("career_direction", "")
    ).strip()
    company_type = guess_company_type(candidate_data, job_intelligence)
    recommended_order = clean_string_list(intelligence.get("priority_sections", []))

    tone = "Business"
    writing_style = "Concise, professional, and ATS-safe."
    resume_strategy = "Balance role alignment, recruiter clarity, and truthful ATS keyword coverage."
    summary_strategy = "Lead with target-role alignment, strongest matching capabilities, and market-ready positioning."
    experience_strategy = "Prioritize role-relevant experience and rewrite evidence with clear impact language."
    project_strategy = "Use projects to demonstrate practical proof of skills when they strengthen the target-role story."
    skills_strategy = "Group skills logically and surface the most role-relevant capabilities first."
    certification_strategy = "Include certifications that strengthen credibility for the target role or market expectations."
    priority_sections = recommended_order[:] or ["Professional Summary", "Skills", "Experience"]
    de_emphasize_sections = []
    industry_language = clean_string_list(ats_intelligence.get("matching_keywords", []))[:6]
    notes = []
    score = 70

    if role_direction.lower() == "technical":
        tone = "Technical"
        writing_style = "Precise, evidence-driven, and technically credible without inflated claims."
        skills_strategy = "Lead with technical skill groups, tools, platforms, and the strongest role-aligned capabilities."
        project_strategy = "Use projects as proof of technical problem-solving, implementation approach, and relevant tool exposure."
        score += 6
    elif role_direction.lower() in {"management", "operations", "sales", "marketing", "finance", "customer support"}:
        tone = "Leadership" if role_direction.lower() in {"management", "operations"} else "Business"
        writing_style = "Commercially clear, outcome-focused, and easy for recruiters and hiring managers to scan."
        experience_strategy = "Emphasize delivery, coordination, communication, process ownership, and business impact."
        score += 5

    if any(token in target_role.lower() for token in ["manager", "head", "director", "lead", "operations manager"]):
        tone = "Executive"
        writing_style = "Authority-led, concise, and focused on leadership, ownership, and business outcomes."
        summary_strategy = "Position the candidate as a leader with operational scope, decision-making value, and strategic relevance."
        experience_strategy = "Highlight team leadership, standards, operational control, and senior-level decision ownership."
        score += 8
    elif experience_level in {"student", "fresher"}:
        tone = "Academic" if tone == "Technical" else tone
        summary_strategy = "Position the candidate through education, projects, internships, certifications, and learning agility."
        priority_sections = clean_string_list(["Professional Summary", "Education", "Projects", "Skills", "Internships", "Certifications"])
        de_emphasize_sections.append("Extended Experience")
        score += 4

    if "startup" in company_type.lower():
        tone = "Startup" if tone not in {"Executive", "Technical"} else tone
        resume_strategy = "Emphasize ownership, adaptability, learning speed, problem solving, and hands-on delivery."
        priority_sections = clean_string_list(priority_sections + ["Projects"])
        de_emphasize_sections.append("Heavy Process Detail")
        notes.append("Startup positioning favors ownership, adaptability, and practical contribution over overly formal wording.")
        score += 7
    elif "enterprise" in company_type.lower() or "corporate" in company_type.lower() or "product company" in company_type.lower():
        resume_strategy = "Emphasize scalable execution, collaboration, documentation, standards, and cross-functional readiness."
        notes.append("Enterprise-style positioning favors clarity, process awareness, and structured execution.")
        score += 6
    elif "government" in company_type.lower():
        resume_strategy = "Emphasize reliability, compliance, communication, procedure adherence, and consistent delivery."
        summary_strategy = "Position the candidate as dependable, process-aware, and aligned to structured environments."
        notes.append("Government-style personalization prioritizes compliance, reliability, and formal clarity.")
        score += 6
    elif "research" in company_type.lower():
        tone = "Academic" if tone != "Executive" else tone
        resume_strategy = "Emphasize experimentation, analytical depth, documentation, and innovation-oriented thinking."
        project_strategy = "Use projects to show experimentation, analysis, tools, methods, and structured findings."
        notes.append("Research positioning favors innovation, experimentation, and analytical rigor.")
        score += 6
    elif "manufacturing" in company_type.lower():
        resume_strategy = "Emphasize process control, reliability, standards, quality, and operational consistency."
        notes.append("Manufacturing positioning should sound structured, practical, and process-aware.")
        score += 5
    elif "hospitality" in company_type.lower() or "retail" in company_type.lower():
        resume_strategy = "Emphasize customer experience, service delivery, coordination, operational consistency, and issue resolution."
        notes.append("Service-sector positioning should highlight coordination, customer handling, and operational consistency.")
        score += 5

    if "united states" in target_country or target_country == "usa":
        writing_style = "Achievement-first, direct, and concise with strong action-led phrasing."
        notes.append("US personalization favors fast value communication and sharp achievement-led wording.")
        score += 5
    elif "canada" in target_country:
        writing_style = "Balanced and professional, combining clear capability statements with grounded evidence."
        notes.append("Canada personalization balances polish, evidence, and readability.")
        score += 4
    elif "united kingdom" in target_country or target_country == "uk":
        writing_style = "Concise, practical, and recruiter-friendly without overstatement."
        notes.append("UK personalization favors concise, practical, and professional wording.")
        score += 4
    elif "australia" in target_country:
        writing_style = "Evidence-based and clear, with practical relevance surfaced early."
        notes.append("Australia personalization favors evidence-based claims and direct relevance.")
        score += 4
    elif "germany" in target_country:
        writing_style = "Structured, precise, and qualification-focused with orderly sectioning."
        notes.append("Germany personalization benefits from precision, structure, and qualifications visibility.")
        score += 5
    elif "uae" in target_country or "united arab emirates" in target_country:
        writing_style = "Polished and leadership-aware, with clear operations and execution language."
        notes.append("UAE personalization favors polished language with leadership and operational clarity.")
        score += 5
    elif "india" in target_country:
        writing_style = "Balanced and practical, giving visible weight to skills, projects, and applied readiness."
        notes.append("India personalization often works best with a balanced emphasis on skills, projects, and readiness.")
        score += 4

    industry_text = " ".join([target_industry.lower(), target_role.lower(), str(job_intelligence.get("industry", "")).lower()])
    if any(token in industry_text for token in ["semiconductor", "vlsi", "electronics", "embedded"]):
        industry_language = clean_string_list(industry_language + ["Design Flow", "Verification", "Linux", "Digital Design", "Implementation"])
        summary_strategy = "Position the candidate through technical credibility, tool exposure, and role-aligned design or verification capability."
        score += 4
    elif any(token in industry_text for token in ["analytics", "business analyst", "data"]):
        industry_language = clean_string_list(industry_language + ["Analysis", "Reporting", "SQL", "Documentation", "Stakeholder Communication"])
        score += 4
    elif any(token in industry_text for token in ["hr", "human resources", "talent"]):
        industry_language = clean_string_list(industry_language + ["Recruitment Support", "Documentation", "Coordination", "Communication", "Onboarding"])
        de_emphasize_sections.append("Deep Technical Detail")
        score += 4
    elif any(token in industry_text for token in ["operations", "supply chain", "restaurant", "hospitality"]):
        industry_language = clean_string_list(industry_language + ["Operations", "Inventory Control", "Service Delivery", "Coordination", "Team Management"])
        score += 4

    if str(getattr(candidate_data, "career_change", "No")).strip().lower() in {"yes", "true", "1"}:
        resume_strategy = "Use a career-switcher narrative that translates prior strengths into target-role value without forcing unrelated detail."
        experience_strategy = "Emphasize transferable achievements, stakeholder impact, and adjacent capability instead of unrelated depth."
        de_emphasize_sections.append("Legacy Domain Detail")
        notes.append("Career-change positioning should translate transferable strengths and reduce unrelated detail.")
        score += 5

    if ats_intelligence.get("missing_keywords"):
        notes.append("Missing ATS keywords should guide future improvement suggestions, not be inserted as unsupported skills.")
        score += 2
    if recruiter_intelligence.get("top_concerns"):
        notes.append("Recruiter concerns should influence emphasis and clarity in the final wording.")
        score += 2

    priority_sections = clean_string_list(priority_sections)[:8]
    de_emphasize_sections = clean_string_list(de_emphasize_sections)[:6]
    recommended_order = clean_string_list(recommended_order or priority_sections)[:8]
    industry_language = clean_string_list(industry_language)[:8]
    overall_note = " ".join(clean_string_list(notes)) or "Tailor the resume tone, emphasis, and section order to the target market and employer context while staying fully truthful."

    return {
        "tone": tone,
        "writing_style": writing_style,
        "resume_strategy": resume_strategy,
        "priority_sections": priority_sections,
        "de_emphasize_sections": de_emphasize_sections,
        "industry_language": industry_language,
        "recommended_order": recommended_order,
        "summary_strategy": summary_strategy,
        "experience_strategy": experience_strategy,
        "project_strategy": project_strategy,
        "skills_strategy": skills_strategy,
        "certification_strategy": certification_strategy,
        "overall_personalization_note": overall_note,
        "personalization_score": max(50, min(98, score)),
        "personalization_strategy": resume_strategy,
        "personalization_notes": clean_string_list(notes)[:8],
    }


def recruiter_review(resume, candidate_data, intelligence):
    system_msg = (
        "You are an experienced recruiter and hiring reviewer with 15+ years of experience. "
        "Review the resume as if you are deciding whether to shortlist the candidate. "
        "Evaluate professional title, summary, experience, projects, skill grouping, ATS alignment, career progression, target role alignment, leadership, achievements, technical depth, transferable skills, and overall readability. "
        "Do not invent facts. Do not expose chain-of-thought. Return only concise structured recruiter feedback."
    )

    user_msg = f"""
Review this resume from the perspective of an experienced recruiter.

Target role: {getattr(candidate_data, 'target_role', '')}
Target country/job market: {getattr(candidate_data, 'target_country', getattr(candidate_data, 'target_location', ''))}
Target industry: {getattr(candidate_data, 'target_industry', '')}
Experience level: {getattr(candidate_data, 'experience_level', '')}
Career direction: {getattr(candidate_data, 'career_direction', '')}

Recruiter context:
Recommended resume model: {intelligence.get('recommended_resume_model', intelligence.get('recommended_resume_style', ''))}
Resume length rule: {intelligence.get('resume_length_rule', '')}
Target market strategy: {intelligence.get('target_market_strategy', '')}
Recruiter positioning: {intelligence.get('recruiter_positioning', '')}
ATS keyword strategy: {json.dumps(intelligence.get('ats_keyword_strategy', []))}
Career graph recommended roles: {json.dumps(intelligence.get('career_graph_roles', []))}
Career graph recommended certifications: {json.dumps(intelligence.get('career_graph_certifications', []))}
Career graph future growth roles: {json.dumps(intelligence.get('career_graph_growth_roles', []))}
Quality score context: {intelligence.get('quality_score', '')}
ATS readiness score context: {intelligence.get('ats_readiness_score', '')}
Recruiter readability score context: {intelligence.get('recruiter_readability_score', '')}
Role alignment score context: {intelligence.get('role_alignment_score', '')}

Resume:
{resume}

Return ONLY valid JSON in this exact format:
{{
    "interview_probability": 0,
    "recruiter_confidence": 0,
    "first_impression": "",
    "shortlisting_decision": "Yes | Maybe | No",
    "top_strengths": ["strength1", "strength2"],
    "top_concerns": ["concern1", "concern2"],
    "missing_high_value_information": ["item1", "item2"],
    "recommended_improvements": ["improvement1", "improvement2"],
    "industry_keywords_missing": ["keyword1", "keyword2"],
    "resume_competitiveness": "Basic | Competitive | Strong | Outstanding"
}}

Scoring rules:
- interview_probability: 0-100
- recruiter_confidence: 0-100
- shortlisting_decision should reflect whether you would shortlist this profile now.
- Be stricter when target role alignment is weak or the resume leaves important uncertainty.
- If the resume feels strong and clear, reward it accordingly.
"""

    resp = client.chat.completions.create(
        model="gpt-4.1-mini",
        messages=[
            {"role": "system", "content": system_msg},
            {"role": "user", "content": user_msg},
        ],
        temperature=0.1,
    )

    content = (resp.choices[0].message.content or "").strip()
    parsed = parse_json_response(content)
    return normalize_recruiter_review(parsed)


def normalize_export_sections(sections: list[ResumeExportSection], country: str) -> list[ResumeExportSection]:
    settings = get_country_template_settings(country)
    normalized = []
    for section in sections:
        heading = section.heading.strip()
        if heading.lower() in {"summary", "professional summary"}:
            heading = settings["summary_heading"]
        elif heading.lower() in {"skills", "key skills"}:
            heading = settings["skills_heading"]

        body = section.body.strip()
        if body:
            normalized.append(ResumeExportSection(heading=heading, body=body))

    return normalized


def build_docx_resume_legacy(export: ResumeExportInput) -> BytesIO:
    settings = get_country_template_settings(export.target_country)
    sections = normalize_export_sections(export.sections, export.target_country)

    document = Document()
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(11)

    title = export.full_name.strip() or export.target_role.strip() or "Resume"
    title_paragraph = document.add_paragraph()
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(18)

    if export.target_role.strip():
      subtitle = document.add_paragraph()
      subtitle_run = subtitle.add_run(export.target_role.strip())
      subtitle_run.italic = True
      subtitle_run.font.size = Pt(11)

    for section in sections:
        heading = document.add_paragraph()
        heading_run = heading.add_run(section.heading)
        heading_run.bold = True
        heading_run.font.size = Pt(12)

        for line in section.body.splitlines():
            cleaned = line.strip()
            if not cleaned:
                continue
            if cleaned.startswith("- ") or cleaned.startswith("* "):
                document.add_paragraph(cleaned[2:].strip(), style="List Bullet")
            else:
                document.add_paragraph(cleaned)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def build_pdf_resume_legacy(export: ResumeExportInput) -> BytesIO:
    settings = get_country_template_settings(export.target_country)
    sections = normalize_export_sections(export.sections, export.target_country)
    buffer = BytesIO()

    doc = SimpleDocTemplate(
        buffer,
        pagesize=settings["page_size"],
        leftMargin=42,
        rightMargin=42,
        topMargin=42,
        bottomMargin=42,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ResumeTitle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=18,
        leading=22,
        alignment=TA_LEFT,
        spaceAfter=8,
    )
    subtitle_style = ParagraphStyle(
        "ResumeSubtitle",
        parent=styles["Normal"],
        fontName="Helvetica-Oblique",
        fontSize=10,
        leading=13,
        textColor="#334155",
        spaceAfter=8,
    )
    heading_style = ParagraphStyle(
        "ResumeHeading",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=11,
        leading=14,
        spaceBefore=8,
        spaceAfter=4,
    )
    body_style = ParagraphStyle(
        "ResumeBody",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=10,
        leading=14,
        spaceAfter=3,
    )

    story = []
    title = export.full_name.strip() or export.target_role.strip() or "Resume"
    story.append(Paragraph(title, title_style))

    if export.target_role.strip():
        story.append(Paragraph(export.target_role.strip(), subtitle_style))

    story.append(Spacer(1, 8))

    for section in sections:
        story.append(Paragraph(section.heading, heading_style))
        for line in section.body.splitlines():
            cleaned = line.strip()
            if not cleaned:
                continue
            if cleaned.startswith("- ") or cleaned.startswith("* "):
                cleaned = f"&bull; {cleaned[2:].strip()}"
            story.append(Paragraph(cleaned.replace("\n", "<br/>"), body_style))
        story.append(Spacer(1, 4))

    doc.build(story)
    buffer.seek(0)
    return buffer


def build_docx_resume(export: ResumeExportInput) -> BytesIO:
    rendered = render_resume_package(export, preferred_theme=export.selected_theme)
    return rendered["docx_buffer"]


def build_pdf_resume(export: ResumeExportInput) -> BytesIO:
    rendered = render_resume_package(export, preferred_theme=export.selected_theme)
    return rendered["pdf_buffer"]


def build_job_alert_preview(data: JobAlertInput) -> tuple[str, str]:
    subject = f"AI Job Copilot Daily Job Alert - {data.target_role}"
    body = f"""Hello,

Your daily AI Job Copilot alert is configured successfully.

Target role: {data.target_role}
Country: {data.country}
City / Location: {data.city}
Experience level: {data.experience_level}
Keywords: {data.keywords or "Not specified"}
Preferred send time: {data.preferred_time}

This is a test/preview email for your job alert setup.

Next production step:
- connect a real job source or API
- run a scheduler daily
- send matched jobs automatically to this inbox

Thanks,
AI Job Copilot
"""
    return subject, body


def send_email(to_email: str, subject: str, body: str) -> None:
    if not SMTP_HOST or not SMTP_USER or not SMTP_PASSWORD or not SMTP_FROM:
        raise RuntimeError("SMTP is not configured.")

    msg = MIMEText(body, "plain", "utf-8")
    msg["Subject"] = subject
    msg["From"] = SMTP_FROM
    msg["To"] = to_email

    with smtplib.SMTP(SMTP_HOST, SMTP_PORT) as server:
        server.starttls()
        server.login(SMTP_USER, SMTP_PASSWORD)
        server.sendmail(SMTP_FROM, [to_email], msg.as_string())


# -----------------------
# Routes
# -----------------------
@app.get("/")
def home():
    return {"status": "running", "service": "ai-job-copilot"}


@app.post("/suggest-role")
def suggest_role(data: RoleSuggestionInput):
    career_knowledge = get_career_knowledge_context(
        data,
        current_background=data.about,
        education=data.about,
        highest_qualification=data.about,
    )
    system_msg = (
        "You are a career assistant for job seekers across technical and non-technical fields. "
        "Read the user's About section and suggest the single best professional role title. "
        "Use the broad role catalog as guidance, but if a better common professional title is needed, "
        "return that role title instead. "
        "Return only one concise role name and nothing else."
    )

    user_msg = f"""
Reference role catalog:
{ROLE_CATALOG}

Career knowledge graph suggestions:
Recommended roles: {career_knowledge['recommended_roles']}
Recommended industries: {career_knowledge['recommended_industries']}
Career transitions: {career_knowledge['career_transition_options']}

User About section:
{data.about}

Choose the single best matching professional role title.
Return only the role name.
"""

    resp = client.chat.completions.create(
        model="gpt-4.1-mini",
        messages=[
            {"role": "system", "content": system_msg},
            {"role": "user", "content": user_msg},
        ],
        temperature=0,
    )

    suggested_role = (resp.choices[0].message.content or "").strip()
    if not suggested_role or len(suggested_role) > 80:
        raise HTTPException(status_code=500, detail=f"Invalid suggested role returned: {suggested_role}")

    return {"suggested_role": suggested_role}


@app.post("/optimize-profile", response_model=OptimizedProfile)
def optimize_profile(data: ProfileInput):
    career_knowledge = get_career_knowledge_context(
        data,
        current_background=data.about,
        education=data.about,
        highest_qualification=data.about,
        target_industry="",
    )
    system_msg = (
        "You are an expert LinkedIn profile strategist for the job market specified by the user. "
        "You optimize profiles for recruiter clarity, keyword relevance, and professional positioning. "
        "The user's About section is the primary source of truth. "
        "You must NEVER invent employers, dates, certifications, achievements, metrics, tools, or projects "
        "that are not clearly supported by the input. "
        "Your job is to improve wording, structure, clarity, and positioning for the target role."
    )

    user_msg = f"""
Target role: {data.target_role}
Target job market / country: {data.target_location}
Career knowledge graph:
Recommended roles: {career_knowledge['recommended_roles']}
Recommended industries: {career_knowledge['recommended_industries']}
Recommended certifications: {career_knowledge['recommended_certifications']}
Recommended projects: {career_knowledge['recommended_projects']}
Recommended resume model: {career_knowledge['recommended_resume_model']}

Primary source input (About section):
{data.about}

Your tasks:
1. Create a LinkedIn headline tailored to the target role.
2. Rewrite the About section in a strong professional tone suitable for the target job market.
3. Keep the About section natural, confident, and recruiter-friendly.
4. Highlight transferable strengths already present in the input.
5. Generate 10-15 relevant LinkedIn/ATS keywords.
6. Do not use emojis.
7. Do not invent facts.
8. If the input is limited, improve presentation but stay truthful.

Headline rules:
- Max 220 characters
- Must be role-first
- Should sound modern and recruiter-friendly
- Avoid generic phrases like "looking for opportunities" unless strongly needed

About rules:
- 120 to 180 words
- Clear opening line
- Mention strengths, tools, and direction only if supported by input
- Make it sound polished but human

Return ONLY valid JSON in this exact format:
{{
  "headline": "string",
  "about": "string",
  "experience_bullets": [],
  "top_keywords": ["keyword1", "keyword2", "keyword3"]
}}
"""

    resp = client.chat.completions.create(
        model="gpt-4.1-mini",
        messages=[
            {"role": "system", "content": system_msg},
            {"role": "user", "content": user_msg},
        ],
        temperature=0.2,
    )

    content = (resp.choices[0].message.content or "").strip()
    parsed = parse_json_response(content)

    required_keys = {"headline", "about", "experience_bullets", "top_keywords"}
    if not required_keys.issubset(parsed.keys()):
        raise HTTPException(
            status_code=500,
            detail=f"Missing keys in model response. Required: {required_keys}. Got: {list(parsed.keys())}",
        )

    if not isinstance(parsed["experience_bullets"], list):
        parsed["experience_bullets"] = []

    if not isinstance(parsed["top_keywords"], list):
        raise HTTPException(status_code=500, detail="top_keywords must be a list.")

    parsed["top_keywords"] = [str(x).strip() for x in parsed["top_keywords"] if str(x).strip()]
    return parsed


@app.post("/generate-profile-from-scratch", response_model=OptimizedProfile)
def generate_profile_from_scratch(data: ScratchProfileInput):
    career_knowledge = get_career_knowledge_context(
        data,
        education=data.education,
        current_background=f"{data.education}\n{data.experience}\n{data.projects}\n{data.career_goal}",
        highest_qualification=data.education,
        target_industry="",
    )
    system_msg = (
        "You are an expert LinkedIn profile strategist for freshers, students, early-career professionals, "
        "and career switchers in the user's target job market. "
        "Your task is to create a strong LinkedIn headline, About section, and keywords from structured user inputs. "
        "You must NEVER invent employers, dates, certifications, achievements, metrics, tools, projects, "
        "or experience that are not clearly supported by the user input. "
        "You should write in a professional, recruiter-friendly style and align the profile to the target role."
    )

    user_msg = f"""
Target role: {data.target_role}
Target job market / country: {data.target_location}
Career knowledge graph:
Recommended roles: {career_knowledge['recommended_roles']}
Recommended industries: {career_knowledge['recommended_industries']}
Recommended certifications: {career_knowledge['recommended_certifications']}
Recommended projects: {career_knowledge['recommended_projects']}
Recommended resume model: {career_knowledge['recommended_resume_model']}

User details:
Education:
{data.education}

Skills:
{data.skills}

Projects:
{data.projects}

Experience:
{data.experience}

Career goal:
{data.career_goal}

Your tasks:
1. Create a LinkedIn headline tailored to the target role.
2. Write a strong LinkedIn About section from scratch using only the provided details.
3. Generate 10-15 relevant LinkedIn/ATS keywords.
4. Do not use emojis.
5. Do not invent facts.
6. If experience is limited, present the profile confidently but truthfully.

Headline rules:
- Max 220 characters
- Must be role-first
- Clear, modern, recruiter-friendly
- Avoid weak phrases like "looking for opportunities" unless needed

About rules:
- 120 to 180 words
- Clear opening line
- Mention education, skills, projects, experience, and goal only if provided
- Sound professional, confident, and natural

Return ONLY valid JSON in this exact format:
{{
  "headline": "string",
  "about": "string",
  "experience_bullets": [],
  "top_keywords": ["keyword1", "keyword2", "keyword3"]
}}
"""

    resp = client.chat.completions.create(
        model="gpt-4.1-mini",
        messages=[
            {"role": "system", "content": system_msg},
            {"role": "user", "content": user_msg},
        ],
        temperature=0.2,
    )

    content = (resp.choices[0].message.content or "").strip()
    parsed = parse_json_response(content)

    required_keys = {"headline", "about", "experience_bullets", "top_keywords"}
    if not required_keys.issubset(parsed.keys()):
        raise HTTPException(
            status_code=500,
            detail=f"Missing keys in model response. Required: {required_keys}. Got: {list(parsed.keys())}",
        )

    if not isinstance(parsed["experience_bullets"], list):
        parsed["experience_bullets"] = []

    if not isinstance(parsed["top_keywords"], list):
        raise HTTPException(status_code=500, detail="top_keywords must be a list.")

    parsed["top_keywords"] = [str(x).strip() for x in parsed["top_keywords"] if str(x).strip()]
    return parsed


@app.post("/optimize-linkedin", response_model=LinkedInOptimizationOutput)
def optimize_linkedin(data: LinkedInOptimizationInput):
    intelligence_request = build_resume_intelligence_request(data)
    intelligence = generate_resume_intelligence(intelligence_request)
    career_knowledge = get_career_knowledge_context(
        data,
        education=data.education_details,
        highest_qualification=data.highest_qualification,
        current_background=data.current_background or data.current_about or data.resume_text,
        target_industry=data.target_industry,
    )
    skill_intelligence = generate_skill_intelligence(data, intelligence)
    job_intelligence = None
    if str(data.job_description or "").strip():
        job_intelligence = analyze_job_description_intelligence(
            JobDescriptionRequest(
                job_title=data.target_role,
                company_name=data.company_name,
                country=data.target_country,
                industry=data.target_industry,
                job_description=data.job_description,
            )
        )
        if job_intelligence.get("recommended_resume_model"):
            intelligence["recommended_resume_model"] = job_intelligence["recommended_resume_model"]
    ats_intelligence = generate_ats_intelligence(
        data,
        resume_text=str(data.resume_text or data.current_about or ""),
        job_intelligence=job_intelligence,
        intelligence=intelligence,
        skill_intelligence=skill_intelligence,
    )
    personalization = generate_resume_personalization(
        data,
        job_intelligence=job_intelligence,
        intelligence=intelligence,
        ats_intelligence=ats_intelligence,
        recruiter_intelligence=None,
    )
    recruiter_source = "\n".join(part for part in [data.resume_text, data.current_about, data.work_experience, data.projects, data.achievements] if str(part or "").strip())
    recruiter_context = dict(intelligence)
    recruiter_context.update({
        "ats_keyword_strategy": ats_intelligence.get("required_keywords", []),
        "target_market_strategy": intelligence.get("target_market_strategy", ""),
        "recruiter_positioning": intelligence.get("recruiter_positioning", ""),
        "career_graph_roles": career_knowledge.get("recommended_roles", []),
        "career_graph_certifications": career_knowledge.get("recommended_certifications", []),
    })
    recruiter_feedback = recruiter_review(recruiter_source, data, recruiter_context)
    return generate_linkedin_optimization_package(
        candidate_data=data,
        intelligence=intelligence,
        skill_intelligence=skill_intelligence,
        job_intelligence=job_intelligence,
        recruiter_intelligence=recruiter_feedback,
        ats_intelligence=ats_intelligence,
        personalization=personalization,
        career_knowledge=career_knowledge,
        client=client,
        parse_json_response=parse_json_response,
    )


@app.post("/generate-interview-prep", response_model=InterviewPrepOutput)
def generate_interview_prep(data: InterviewPrepInput):
    intelligence_request = build_resume_intelligence_request(data)
    intelligence = generate_resume_intelligence(intelligence_request)
    career_knowledge = get_career_knowledge_context(
        data,
        education=data.education_details,
        highest_qualification=data.highest_qualification,
        current_background=data.current_background or data.resume_text,
        target_industry=data.target_industry,
    )
    skill_intelligence = generate_skill_intelligence(data, intelligence)
    job_intelligence = None
    if str(data.job_description or "").strip():
        job_intelligence = analyze_job_description_intelligence(
            JobDescriptionRequest(
                job_title=data.target_role,
                company_name=data.company_name,
                country=data.target_country,
                industry=data.target_industry,
                job_description=data.job_description,
            )
        )
        if job_intelligence.get("recommended_resume_model"):
            intelligence["recommended_resume_model"] = job_intelligence["recommended_resume_model"]
    ats_intelligence = generate_ats_intelligence(
        data,
        resume_text=str(data.resume_text or data.current_background or ""),
        job_intelligence=job_intelligence,
        intelligence=intelligence,
        skill_intelligence=skill_intelligence,
    )
    recruiter_source = "\n".join(part for part in [data.resume_text, data.current_background, data.work_experience, data.internships, data.projects, data.achievements] if str(part or "").strip())
    recruiter_context = dict(intelligence)
    recruiter_context.update({
        "ats_keyword_strategy": ats_intelligence.get("required_keywords", []),
        "target_market_strategy": intelligence.get("target_market_strategy", ""),
        "recruiter_positioning": intelligence.get("recruiter_positioning", ""),
        "career_graph_roles": career_knowledge.get("recommended_roles", []),
        "career_graph_certifications": career_knowledge.get("recommended_certifications", []),
        "career_graph_growth_roles": career_knowledge.get("future_growth_roles", []),
    })
    recruiter_feedback = recruiter_review(recruiter_source, data, recruiter_context)
    return generate_interview_prep_package(
        candidate_data=data,
        intelligence=intelligence,
        job_intelligence=job_intelligence,
        recruiter_intelligence=recruiter_feedback,
        ats_intelligence=ats_intelligence,
        career_knowledge=career_knowledge,
        client=client,
        parse_json_response=parse_json_response,
    )


@app.post("/generate-portfolio", response_model=PortfolioOutput)
def generate_portfolio(data: PortfolioInput):
    intelligence_request = build_resume_intelligence_request(data)
    intelligence = generate_resume_intelligence(intelligence_request)
    career_knowledge = get_career_knowledge_context(
        data,
        education=data.education_details,
        highest_qualification=data.highest_qualification,
        current_background=data.current_background or data.resume_text,
        target_industry=data.target_industry,
    )
    skill_intelligence = generate_skill_intelligence(data, intelligence)
    job_intelligence = None
    if str(data.job_description or "").strip():
        job_intelligence = analyze_job_description_intelligence(
            JobDescriptionRequest(
                job_title=data.target_role,
                company_name=data.company_name,
                country=data.target_country,
                industry=data.target_industry,
                job_description=data.job_description,
            )
        )
        if job_intelligence.get("recommended_resume_model"):
            intelligence["recommended_resume_model"] = job_intelligence["recommended_resume_model"]
    achievement_intelligence = generate_achievement_intelligence(data, intelligence, job_intelligence)
    ats_intelligence = generate_ats_intelligence(
        data,
        resume_text=str(data.resume_text or data.current_background or ""),
        job_intelligence=job_intelligence,
        intelligence=intelligence,
        skill_intelligence=skill_intelligence,
    )
    recruiter_source = "\n".join(part for part in [data.resume_text, data.current_background, data.work_experience, data.internships, data.projects, data.achievements] if str(part or "").strip())
    recruiter_context = dict(intelligence)
    recruiter_context.update({
        "ats_keyword_strategy": ats_intelligence.get("required_keywords", []),
        "target_market_strategy": intelligence.get("target_market_strategy", ""),
        "recruiter_positioning": intelligence.get("recruiter_positioning", ""),
        "career_graph_roles": career_knowledge.get("recommended_roles", []),
        "career_graph_certifications": career_knowledge.get("recommended_certifications", []),
        "career_graph_growth_roles": career_knowledge.get("future_growth_roles", []),
    })
    recruiter_feedback = recruiter_review(recruiter_source, data, recruiter_context)
    linkedin_context = generate_linkedin_optimization_package(
        candidate_data=data,
        intelligence=intelligence,
        skill_intelligence=skill_intelligence,
        job_intelligence=job_intelligence,
        recruiter_intelligence=recruiter_feedback,
        ats_intelligence=ats_intelligence,
        personalization=generate_resume_personalization(data, job_intelligence=job_intelligence, intelligence=intelligence, ats_intelligence=ats_intelligence, recruiter_intelligence=None),
        career_knowledge=career_knowledge,
        client=client,
        parse_json_response=parse_json_response,
    )
    interview_context = generate_interview_prep_package(
        candidate_data=data,
        intelligence=intelligence,
        job_intelligence=job_intelligence,
        recruiter_intelligence=recruiter_feedback,
        ats_intelligence=ats_intelligence,
        career_knowledge=career_knowledge,
        client=client,
        parse_json_response=parse_json_response,
    )
    return generate_portfolio_package(
        candidate_data=data,
        intelligence=intelligence,
        skill_intelligence=skill_intelligence,
        achievement_intelligence=achievement_intelligence,
        recruiter_intelligence=recruiter_feedback,
        ats_intelligence=ats_intelligence,
        career_knowledge=career_knowledge,
        linkedin_context=linkedin_context,
        interview_context=interview_context,
        job_intelligence=job_intelligence,
        client=client,
        parse_json_response=parse_json_response,
    )


@app.post("/analyze-resume-intelligence", response_model=ResumeIntelligenceAnalysisOutput)
def analyze_resume_intelligence(data: ResumeIntelligenceRequest):
    return generate_resume_intelligence(data)


@app.post("/analyze-job-description", response_model=JobDescriptionAnalysisOutput)
def analyze_job_description(data: JobDescriptionRequest):
    return analyze_job_description_intelligence(data)


@app.post("/analyze-career-knowledge", response_model=CareerKnowledgeOutput)
def analyze_career_knowledge(data: CareerKnowledgeRequest):
    result = generate_career_knowledge(data.model_dump())
    return {
        "recommended_roles": result["recommended_roles"],
        "recommended_industries": result["recommended_industries"],
        "recommended_certifications": result["recommended_certifications"],
        "recommended_projects": result["recommended_projects"],
        "recommended_resume_model": result["recommended_resume_model"],
        "career_transition_options": result["career_transition_options"],
        "future_growth_roles": result["future_growth_roles"],
        "salary_progression_note": result["salary_progression_note"],
    }


@app.post("/analyze-ats", response_model=ATSIntelligenceOutput)
def analyze_ats(data: ATSAnalysisRequest):
    job_intelligence = None
    if str(data.job_description or "").strip():
        job_intelligence = analyze_job_description_intelligence(
            JobDescriptionRequest(
                job_title=data.target_role,
                company_name="",
                country=data.target_country,
                industry=data.target_industry,
                job_description=data.job_description,
            )
        )
    skill_intelligence = generate_skill_intelligence(data, intelligence=None)
    return generate_ats_intelligence(
        data,
        resume_text=data.resume_text,
        job_intelligence=job_intelligence,
        intelligence=None,
        skill_intelligence=skill_intelligence,
    )


@app.post("/generate-achievements", response_model=AchievementIntelligenceOutput)
def generate_achievements(data: AchievementRequest):
    return generate_achievement_intelligence(data)


@app.post("/optimize-resume", response_model=ResumeOptimizerOutput)
def optimize_resume(data: ResumeOptimizerInput):
    country_rules = get_country_rules(data.target_country)
    career_knowledge = get_career_knowledge_context(
        data,
        education=data.resume_text,
        highest_qualification=data.resume_text,
        current_background=data.resume_text,
        target_industry="",
    )
    job_intelligence = None
    skill_match = {"skill_match_percentage": 0, "strong_matching_skills": [], "missing_required_skills": [], "missing_preferred_skills": [], "resume_alignment_strategy": ""}
    if str(data.job_description or "").strip():
        job_intelligence = analyze_job_description_intelligence(JobDescriptionRequest(job_title=data.target_role, company_name="", country=data.target_country, industry="", job_description=data.job_description))
        skill_match = compare_candidate_to_job_description(data, job_intelligence)
    ats_intelligence = generate_ats_intelligence(
        ATSAnalysisRequest(
            target_role=data.target_role,
            target_country=data.target_country,
            target_industry="",
            experience_level="",
            career_direction="",
            resume_text=data.resume_text,
            job_description=data.job_description,
            technical_skills="",
            transferable_skills="",
            tools_software="",
            projects=data.resume_text,
            work_experience=data.resume_text,
        ),
        resume_text=data.resume_text,
        job_intelligence=job_intelligence,
        intelligence=None,
        skill_intelligence=None,
    )
    personalization = generate_resume_personalization(
        data,
        job_intelligence=job_intelligence,
        intelligence=None,
        ats_intelligence=ats_intelligence,
        recruiter_intelligence=None,
    )
    achievement_intelligence = generate_achievement_intelligence(
        AchievementRequest(
            target_role=data.target_role,
            target_country=data.target_country,
            experience_level="",
            career_direction="",
            work_experience=data.resume_text,
            internships="",
            projects=data.resume_text,
            achievements="",
            leadership_experience="",
            technical_skills="",
            transferable_skills="",
            resume_text=data.resume_text,
        ),
        intelligence=None,
        job_intelligence=job_intelligence,
    )
    system_msg = (
        "You are a senior professional resume writer, ATS optimization specialist, and resume rebuilder with 15 years of experience. "
        "Your job is to analyze an existing resume and rebuild it into a stronger, recruiter-quality, ATS-friendly resume. "
        "You must think like a recruiter, hiring manager, ATS parser, and career strategist before writing. "
        "You must NEVER invent employers, dates, certifications, achievements, metrics, tools, projects, or responsibilities not supported by the user's resume text. "
        "If a job description is provided, align the wording and keyword emphasis to it, but remain truthful. "
        "Adapt tone and structure to the target role, target country, and likely seniority level suggested by the input. "
        "Do not create generic AI summaries. Return a complete optimized resume, not just fragments. "
        "Never include labels like Template: India, Template: UK, preview, backend, or test in the output. "
        "Use clean ATS-safe structure, strong action verbs, clear role alignment, and keyword relevance. "
        "Avoid tables, columns, icons, graphics, or decorative formatting guidance."
    )

    user_msg = f"""
Target role: {data.target_role}
Target country / job market: {data.target_country}

Resume text:
{data.resume_text}

Optional Job Description:
{data.job_description}

Job Description Intelligence:
{json.dumps(job_intelligence) if job_intelligence else "Not provided"}
Skill Match Intelligence:
{json.dumps(skill_match)}
ATS Intelligence:
{json.dumps(ats_intelligence)}
Career Knowledge Graph:
{json.dumps(career_knowledge)}
Resume Personalization:
{json.dumps(personalization)}
Achievement Intelligence:
{json.dumps(achievement_intelligence)}

Country Rules Engine:
{country_rules}

Resume Style Engine:
{get_resume_style_models()}

Tasks:
1. Analyze the current resume and identify the biggest weaknesses blocking recruiter impact or ATS performance.
2. Recommend the best resume style for this candidate based on the content, target role, target country, and likely experience level inferred from the resume.
3. Explain briefly why that style is suitable.
4. Rewrite the entire resume professionally into a complete ATS-friendly version.
5. Tailor the wording to the target role and job market.
6. If a job description is provided, prioritize matching terminology where truthful.
7. Group skills and competencies in a cleaner, more recruiter-friendly structure.
8. Strengthen weak bullet points into clearer impact-oriented statements only when supported by the user's content.
9. Reuse the Achievement Intelligence to rebuild weak responsibilities, projects, and experience bullets into stronger truthful statements.
10. Do not invent facts.
11. Do not use emojis.
12. Do not include Template labels or developer/testing language.
13. Group skills into logical clusters instead of listing them as a flat, repetitive skills block.
14. Keep length discipline appropriate to the likely seniority level suggested by the resume.
15. Use the ATS Intelligence placement strategy to improve keyword coverage naturally without keyword stuffing.
16. Missing keywords may guide suggestions, but must never be added as fake skills or unsupported experience.
17. Follow the Resume Personalization strategy so the wording, emphasis, and structure match the target country, company type, and industry context.
18. Use the Career Knowledge Graph to keep role language, certifications, projects, and growth positioning realistic and consistent.

Resume rebuild rules:
- Include a clear header/contact section if details exist in the pasted resume
- Include a professional title aligned to the target role when supported
- Include a professional summary
- Include key skills or core competencies
- Include professional experience / relevant experience
- Include projects if relevant content exists
- Include education if present
- Include certifications if present
- Keep it ATS-friendly and paste-ready

Resume length rules:
- Student / Fresher: maximum 1 page
- 1-5 years: 1 page
- 5-10 years: 1-2 pages
- 10+ years: 2 pages maximum
- Prioritize relevance and avoid repetition

ATS score estimate rules:
- Provide a short estimate like "High ATS readiness (82-88/100 range)" or "Moderate ATS readiness (68-75/100 range)"
- This is an internal estimate, not a guaranteed external ATS score

Return ONLY valid JSON in this exact format:
{{
  "recommended_resume_style": "string",
  "recommendation_reason": "string",
  "optimized_resume": "string",
  "ats_keywords": ["keyword1", "keyword2", "keyword3"],
  "strengths": ["strength1", "strength2"],
  "weaknesses_found": ["weakness1", "weakness2"],
  "improvement_suggestions": ["suggestion1", "suggestion2"],
  "ats_score_estimate": 0,
  "ats_readiness_level": "string",
  "matching_keywords": ["keyword1", "keyword2"],
  "missing_keywords": ["keyword1", "keyword2"],
  "ats_improvement_actions": ["action1", "action2"],
  "personalization_score": 0,
  "personalization_strategy": "string",
  "personalization_notes": ["note1", "note2"]
}}
"""

    resp = client.chat.completions.create(
        model="gpt-4.1-mini",
        messages=[
            {"role": "system", "content": system_msg},
            {"role": "user", "content": user_msg},
        ],
        temperature=0.2,
    )

    content = (resp.choices[0].message.content or "").strip()
    parsed = parse_json_response(content)

    required_keys = {
        "recommended_resume_style",
        "recommendation_reason",
        "optimized_resume",
        "ats_keywords",
        "strengths",
        "weaknesses_found",
        "improvement_suggestions",
        "ats_score_estimate",
        "ats_readiness_level",
        "matching_keywords",
        "missing_keywords",
        "ats_improvement_actions",
        "personalization_score",
        "personalization_strategy",
        "personalization_notes",
    }
    if not required_keys.issubset(parsed.keys()):
        raise HTTPException(
            status_code=500,
            detail=f"Missing keys in model response. Required: {required_keys}. Got: {list(parsed.keys())}",
        )

    if not isinstance(parsed["ats_keywords"], list):
        parsed["ats_keywords"] = []
    if not isinstance(parsed["strengths"], list):
        parsed["strengths"] = []
    if not isinstance(parsed["weaknesses_found"], list):
        parsed["weaknesses_found"] = []
    if not isinstance(parsed["improvement_suggestions"], list):
        parsed["improvement_suggestions"] = []
    if not isinstance(parsed["matching_keywords"], list):
        parsed["matching_keywords"] = []
    if not isinstance(parsed["missing_keywords"], list):
        parsed["missing_keywords"] = []
    if not isinstance(parsed["ats_improvement_actions"], list):
        parsed["ats_improvement_actions"] = []
    if not isinstance(parsed["personalization_notes"], list):
        parsed["personalization_notes"] = []

    parsed["optimized_resume"] = sanitize_resume_text(str(parsed.get("optimized_resume", "")).strip())
    try:
        parsed["ats_score_estimate"] = int(parsed.get("ats_score_estimate", 0))
    except Exception:
        parsed["ats_score_estimate"] = 0
    try:
        parsed["personalization_score"] = int(parsed.get("personalization_score", 0))
    except Exception:
        parsed["personalization_score"] = 0
    parsed["ats_readiness_level"] = str(parsed.get("ats_readiness_level", "")).strip()
    parsed["personalization_strategy"] = str(parsed.get("personalization_strategy", "")).strip()
    parsed["ats_keywords"] = [str(x).strip() for x in parsed["ats_keywords"] if str(x).strip()]
    parsed["strengths"] = [str(x).strip() for x in parsed["strengths"] if str(x).strip()]
    parsed["weaknesses_found"] = [str(x).strip() for x in parsed["weaknesses_found"] if str(x).strip()]
    parsed["improvement_suggestions"] = [str(x).strip() for x in parsed["improvement_suggestions"] if str(x).strip()]
    parsed["matching_keywords"] = [str(x).strip() for x in parsed["matching_keywords"] if str(x).strip()]
    parsed["missing_keywords"] = [str(x).strip() for x in parsed["missing_keywords"] if str(x).strip()]
    parsed["ats_improvement_actions"] = [str(x).strip() for x in parsed["ats_improvement_actions"] if str(x).strip()]
    parsed["personalization_notes"] = [str(x).strip() for x in parsed["personalization_notes"] if str(x).strip()]

    final_ats = generate_ats_intelligence(
        ATSAnalysisRequest(
            target_role=data.target_role,
            target_country=data.target_country,
            target_industry="",
            experience_level="",
            career_direction="",
            resume_text=parsed["optimized_resume"],
            job_description=data.job_description,
            technical_skills="",
            transferable_skills="",
            tools_software="",
            projects=parsed["optimized_resume"],
            work_experience=parsed["optimized_resume"],
        ),
        resume_text=parsed["optimized_resume"],
        job_intelligence=job_intelligence,
        intelligence=None,
        skill_intelligence=None,
    )
    parsed["ats_score_estimate"] = final_ats["ats_score_estimate"]
    parsed["ats_readiness_level"] = final_ats["ats_readiness_level"]
    parsed["matching_keywords"] = final_ats["matching_keywords"]
    parsed["missing_keywords"] = final_ats["missing_keywords"]
    parsed["ats_improvement_actions"] = final_ats["ats_improvement_actions"]
    parsed["ats_keywords"] = clean_string_list(parsed["ats_keywords"] or final_ats["required_keywords"])
    parsed["improvement_suggestions"] = clean_string_list(parsed["improvement_suggestions"] + final_ats["ats_improvement_actions"])
    final_personalization = generate_resume_personalization(
        data,
        job_intelligence=job_intelligence,
        intelligence=None,
        ats_intelligence=final_ats,
        recruiter_intelligence=None,
    )
    parsed["personalization_score"] = final_personalization["personalization_score"]
    parsed["personalization_strategy"] = final_personalization["personalization_strategy"]
    parsed["personalization_notes"] = final_personalization["personalization_notes"]

    recruiter_context = {
        "recommended_resume_style": parsed.get("recommended_resume_style", ""),
        "resume_length_rule": "",
        "target_market_strategy": country_rules,
        "recruiter_positioning": parsed.get("recommendation_reason", ""),
        "ats_keyword_strategy": parsed.get("ats_keywords", []),
        "career_graph_roles": career_knowledge["recommended_roles"],
        "career_graph_certifications": career_knowledge["recommended_certifications"],
        "career_graph_growth_roles": career_knowledge["future_growth_roles"],
    }
    recruiter_feedback = recruiter_review(parsed["optimized_resume"], data, recruiter_context)
    parsed.update(recruiter_feedback)
    parsed.update(skill_match)
    parsed["personalization_notes"] = clean_string_list(
        parsed["personalization_notes"] + [f"Overall: {final_personalization['overall_personalization_note']}"]
    )
    optimized_export = ResumeExportInput(
        full_name="",
        target_role=data.target_role,
        target_country=data.target_country,
        sections=resume_text_to_sections(parsed["optimized_resume"]),
        selected_theme="",
        experience_level="",
    )
    rendered = render_resume_package(
        optimized_export,
        resume_model=parsed["recommended_resume_style"],
        preferred_theme="",
    )
    parsed["resume_pdf_path"] = rendered["resume_pdf_path"]
    parsed["resume_docx_path"] = rendered["resume_docx_path"]
    parsed["selected_theme"] = rendered["selected_theme"]
    parsed["page_count"] = rendered["page_count"]
    parsed["render_quality_score"] = rendered["render_quality_score"]
    return parsed


@app.post("/build-resume", response_model=ResumeBuildOutput)
def build_resume(data: ResumeIntelligenceInput):
    intelligence_request = build_resume_intelligence_request(data)
    intelligence = generate_resume_intelligence(intelligence_request)
    career_knowledge = get_career_knowledge_context(
        data,
        education=data.education_details,
        highest_qualification=data.highest_qualification,
        current_background=data.current_background,
        target_industry=data.target_industry,
    )
    skill_intelligence = generate_skill_intelligence(data, intelligence)
    job_intelligence = None
    skill_match = {"skill_match_percentage": 0, "strong_matching_skills": [], "missing_required_skills": [], "missing_preferred_skills": [], "resume_alignment_strategy": ""}
    if str(data.job_description or "").strip():
        job_intelligence = analyze_job_description_intelligence(JobDescriptionRequest(job_title=data.target_role, company_name=data.company_name, country=data.target_country, industry=data.target_industry, job_description=data.job_description))
        skill_match = compare_candidate_to_job_description(data, job_intelligence)
        if job_intelligence.get("recommended_resume_model"):
            intelligence["recommended_resume_model"] = job_intelligence["recommended_resume_model"]
    ats_intelligence = generate_ats_intelligence(
        data,
        resume_text="",
        job_intelligence=job_intelligence,
        intelligence=intelligence,
        skill_intelligence=skill_intelligence,
    )
    personalization = generate_resume_personalization(
        data,
        job_intelligence=job_intelligence,
        intelligence=intelligence,
        ats_intelligence=ats_intelligence,
        recruiter_intelligence=None,
    )
    achievement_intelligence = generate_achievement_intelligence(data, intelligence, job_intelligence)
    intelligence["career_graph_roles"] = career_knowledge["recommended_roles"]
    intelligence["career_graph_certifications"] = career_knowledge["recommended_certifications"]
    intelligence["career_graph_growth_roles"] = career_knowledge["future_growth_roles"]
    intelligence["target_role"] = data.target_role
    intelligence["experience_level"] = data.experience_level
    selected_resume_style, resume_model_module = select_resume_model(intelligence)
    intelligence["recommended_resume_model"] = selected_resume_style

    recommendation_reason = (
        f"Selected {selected_resume_style} based on the candidate profile type '{intelligence.get('candidate_profile_type', '')}', "
        f"career direction '{intelligence.get('career_direction_detected', '')}', and target role alignment for {data.target_role}."
    )

    def build_strengths() -> list[str]:
        strengths = []
        strengths.extend(skill_intelligence.get("priority_skills", [])[:3])
        strengths.extend(item.get("improved", "") for item in achievement_intelligence.get("experience_bullets", [])[:2])
        strengths.extend(item.get("project_value", "") for item in achievement_intelligence.get("project_bullets", [])[:1])
        return clean_string_list(strengths)[:6]

    def generate_resume_payload(rewrite_context: dict | None = None) -> dict:
        runtime_intelligence = dict(intelligence)
        if rewrite_context:
            runtime_intelligence["rewrite_context"] = rewrite_context
        model_output = resume_model_module.generate_resume(
            data,
            runtime_intelligence,
            skill_intelligence,
            achievement_intelligence,
            ats_intelligence,
            rewrite_context or {},
            job_intelligence,
            personalization,
        )
        return {
            "recommended_resume_style": selected_resume_style,
            "recommendation_reason": recommendation_reason,
            "professional_title": model_output.get("professional_title", ""),
            "executive_summary": model_output.get("summary", ""),
            "resume_length_rule": intelligence["resume_length_rule"],
            "target_market_strategy": intelligence["target_market_strategy"],
            "recruiter_positioning": intelligence["recruiter_positioning"],
            "full_resume": model_output.get("resume", ""),
            "ats_keywords": ats_intelligence.get("required_keywords", []),
            "skill_groups": skill_intelligence.get("skill_groups", []),
            "strengths": build_strengths(),
            "missing_information": intelligence.get("missing_information", []),
            "improvement_suggestions": [],
            "writing_quality_score": "",
            "resume_readability": "",
            "ats_readiness": ats_intelligence.get("ats_readiness_level", ""),
            "resume_confidence": "",
        }

    parsed = generate_resume_payload()
    parsed["full_resume"] = sanitize_resume_text(str(parsed.get("full_resume", "")).strip())
    quality_report = review_resume_quality(parsed["full_resume"], data, intelligence, skill_intelligence)
    quality_fixes_applied = []

    if (not quality_report["is_ready_for_user"]) or quality_report["quality_score"] < 80:
        quality_fixes_applied = clean_string_list(quality_report["required_fixes"])
        parsed = generate_resume_payload(
            {
                "draft_resume": parsed["full_resume"],
                "draft_summary": parsed.get("executive_summary", ""),
                "issues_found": quality_report["issues_found"],
                "required_fixes": quality_report["required_fixes"],
            }
        )
        parsed["full_resume"] = sanitize_resume_text(str(parsed.get("full_resume", "")).strip())
        quality_report = review_resume_quality(parsed["full_resume"], data, intelligence, skill_intelligence)

    final_response = normalize_build_resume_response(parsed, intelligence, skill_intelligence, quality_report, quality_fixes_applied)
    final_ats = generate_ats_intelligence(
        data,
        resume_text=final_response["full_resume"],
        job_intelligence=job_intelligence,
        intelligence=intelligence,
        skill_intelligence=skill_intelligence,
    )
    final_response["ats_score_estimate"] = final_ats["ats_score_estimate"]
    final_response["ats_readiness_level"] = final_ats["ats_readiness_level"]
    final_response["matching_keywords"] = final_ats["matching_keywords"]
    final_response["missing_keywords"] = final_ats["missing_keywords"]
    final_response["ats_improvement_actions"] = final_ats["ats_improvement_actions"]
    final_response["improvement_suggestions"] = clean_string_list(final_response["improvement_suggestions"] + final_ats["ats_improvement_actions"])
    final_personalization = generate_resume_personalization(
        data,
        job_intelligence=job_intelligence,
        intelligence=intelligence,
        ats_intelligence=final_ats,
        recruiter_intelligence=None,
    )
    final_response["personalization_score"] = final_personalization["personalization_score"]
    final_response["personalization_strategy"] = final_personalization["personalization_strategy"]
    final_response["personalization_notes"] = clean_string_list(
        final_personalization["personalization_notes"] + [f"Overall: {final_personalization['overall_personalization_note']}"]
    )
    final_response.update(skill_match)
    recruiter_context = dict(intelligence)
    recruiter_context.update({
        "quality_score": quality_report["quality_score"],
        "ats_readiness_score": quality_report["ats_readiness_score"],
        "recruiter_readability_score": quality_report["recruiter_readability_score"],
        "role_alignment_score": quality_report["role_alignment_score"],
        "career_graph_roles": career_knowledge["recommended_roles"],
        "career_graph_certifications": career_knowledge["recommended_certifications"],
        "career_graph_growth_roles": career_knowledge["future_growth_roles"],
    })
    recruiter_feedback = recruiter_review(final_response["full_resume"], data, recruiter_context)
    final_response.update(recruiter_feedback)
    rendered = render_resume_package(
        build_export_input_from_resume_response(final_response, data, "full_resume"),
        resume_model=final_response["recommended_resume_style"],
        preferred_theme="",
    )
    final_response["resume_pdf_path"] = rendered["resume_pdf_path"]
    final_response["resume_docx_path"] = rendered["resume_docx_path"]
    final_response["selected_theme"] = rendered["selected_theme"]
    final_response["page_count"] = rendered["page_count"]
    final_response["render_quality_score"] = rendered["render_quality_score"]
    return final_response


@app.post("/generate-cover-letter", response_model=CoverLetterOutput)
def generate_cover_letter_endpoint(data: CoverLetterInput):
    intelligence_request = build_resume_intelligence_request(data)
    intelligence = generate_resume_intelligence(intelligence_request)
    career_knowledge = get_career_knowledge_context(
        data,
        education=data.education_details,
        highest_qualification=data.highest_qualification,
        current_background=data.current_background,
        target_industry=data.target_industry,
    )
    skill_intelligence = generate_skill_intelligence(data, intelligence)
    job_intelligence = None
    if str(data.job_description or "").strip():
        job_intelligence = analyze_job_description_intelligence(
            JobDescriptionRequest(
                job_title=data.target_role,
                company_name=data.company_name,
                country=data.target_country,
                industry=data.target_industry,
                job_description=data.job_description,
            )
        )
        if job_intelligence.get("recommended_resume_model"):
            intelligence["recommended_resume_model"] = job_intelligence["recommended_resume_model"]

    achievement_intelligence = generate_achievement_intelligence(data, intelligence, job_intelligence)
    ats_intelligence = generate_ats_intelligence(
        data,
        resume_text=str(data.resume_text or ""),
        job_intelligence=job_intelligence,
        intelligence=intelligence,
        skill_intelligence=skill_intelligence,
    )
    personalization = generate_resume_personalization(
        data,
        job_intelligence=job_intelligence,
        intelligence=intelligence,
        ats_intelligence=ats_intelligence,
        recruiter_intelligence=None,
    )
    selected_resume_style, _ = select_resume_model(intelligence)
    intelligence["recommended_resume_model"] = selected_resume_style
    intelligence["career_graph_roles"] = career_knowledge["recommended_roles"]
    intelligence["career_graph_certifications"] = career_knowledge["recommended_certifications"]
    intelligence["career_graph_growth_roles"] = career_knowledge["future_growth_roles"]

    reference_resume_text = str(data.resume_text or "").strip()
    if not reference_resume_text:
        grouped_skills = []
        for group in skill_intelligence.get("skill_groups", []):
            grouped_skills.extend(group.get("skills", []))
        reference_resume_text = "\n".join(
            part
            for part in [
                data.full_name,
                data.target_role,
                "Professional Summary",
                data.current_background,
                "Work Experience",
                data.work_experience,
                "Internships",
                data.internships,
                "Projects",
                data.projects,
                "Skills",
                ", ".join(clean_string_list(grouped_skills)),
                "Achievements",
                "\n".join(item.get("improved", "") for item in achievement_intelligence.get("experience_bullets", [])[:3]),
            ]
            if str(part or "").strip()
        )

    recruiter_context = dict(intelligence)
    recruiter_context.update({
        "ats_keyword_strategy": ats_intelligence.get("required_keywords", []),
        "target_market_strategy": intelligence.get("target_market_strategy", ""),
        "recruiter_positioning": intelligence.get("recruiter_positioning", ""),
    })
    recruiter_feedback = recruiter_review(reference_resume_text, data, recruiter_context)

    return generate_cover_letter_package(
        candidate_data=data,
        intelligence=intelligence,
        skill_intelligence=skill_intelligence,
        achievement_intelligence=achievement_intelligence,
        job_intelligence=job_intelligence,
        recruiter_intelligence=recruiter_feedback,
        ats_intelligence=ats_intelligence,
        personalization=personalization,
        career_knowledge=career_knowledge,
        resume_model_name=selected_resume_style,
        client=client,
        parse_json_response=parse_json_response,
    )


@app.post("/review-resume", response_model=ResumeReviewerOutput)
def review_resume(data: ResumeReviewerInput):
    country_rules = get_country_rules(data.target_country)
    system_msg = (
        "You are a senior ATS optimization specialist, recruiter, and resume reviewer with 15 years of experience. "
        "Your job is to review an existing resume like a recruiter and ATS consultant would. "
        "You must identify what is working, what is weak, which keywords are missing, and why the ATS estimate was given. "
        "Do not invent facts from the resume. "
        "Do not use generic AI wording. "
        "Do not include template labels, preview, backend, or test language."
    )

    user_msg = f"""
Review this resume like a professional resume reviewer.

Target role: {data.target_role}
Target country / job market: {data.target_country}
Optional job description:
{data.job_description}

Country Rules Engine:
{country_rules}

Resume text:
{data.resume_text}

Tasks:
1. Estimate ATS readiness.
2. Explain clearly why that score estimate was given.
3. Identify missing keywords for the target role when possible.
4. Identify weak sections.
5. Identify strong sections.
6. Write recruiter feedback in a concise professional tone.
7. Give practical improvement suggestions.
8. Do not invent missing experiences, metrics, or employers.

ATS score format rules:
- Return ATS score in a professional estimate style such as "Low ATS readiness (30-40/100 range)" or "Moderate ATS readiness (60-70/100 range)"
- Do not return a raw percentage only

Return ONLY valid JSON in this exact format:
{{
  "ats_score_estimate": "string",
  "score_reason": "string",
  "missing_keywords": ["keyword1", "keyword2"],
  "weak_sections": ["section1", "section2"],
  "strong_sections": ["section1", "section2"],
  "recruiter_feedback": "string",
  "improvement_suggestions": ["suggestion1", "suggestion2"]
}}
"""

    resp = client.chat.completions.create(
        model="gpt-4.1-mini",
        messages=[
            {"role": "system", "content": system_msg},
            {"role": "user", "content": user_msg},
        ],
        temperature=0.2,
    )

    content = (resp.choices[0].message.content or "").strip()
    parsed = parse_json_response(content)

    required_keys = {
        "ats_score_estimate",
        "score_reason",
        "missing_keywords",
        "weak_sections",
        "strong_sections",
        "recruiter_feedback",
        "improvement_suggestions",
    }
    if not required_keys.issubset(parsed.keys()):
        raise HTTPException(
            status_code=500,
            detail=f"Missing keys in resume review response. Required: {required_keys}. Got: {list(parsed.keys())}",
        )

    for key in ("missing_keywords", "weak_sections", "strong_sections", "improvement_suggestions"):
        if not isinstance(parsed[key], list):
            parsed[key] = []
        parsed[key] = [str(x).strip() for x in parsed[key] if str(x).strip()]

    return parsed


@app.post("/export-resume-docx")
def export_resume_docx(data: ResumeExportInput):
    buffer = build_docx_resume(data)
    filename = f"{(data.full_name or data.target_role or 'resume').strip().replace(' ', '-').lower()}-resume.docx"
    headers = {"Content-Disposition": f'attachment; filename="{filename}"'}
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
    )


@app.post("/export-resume-pdf")
def export_resume_pdf(data: ResumeExportInput):
    buffer = build_pdf_resume(data)
    filename = f"{(data.full_name or data.target_role or 'resume').strip().replace(' ', '-').lower()}-resume.pdf"
    headers = {"Content-Disposition": f'attachment; filename="{filename}"'}
    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers=headers,
    )


@app.post("/generate-hiring-messages", response_model=HiringMessageOutput)
def generate_hiring_messages(data: HiringMessageInput):
    system_msg = (
        "You are an expert career communication assistant for the user's target job market. "
        "Your job is to generate short, polite, professional outreach messages for LinkedIn or email. "
        "You must NEVER invent achievements, years of experience, or technical skills not supported by the user's background. "
        "Messages should feel human, confident, and respectful—not spammy or over-salesy."
    )

    user_msg = f"""
Target role: {data.target_role}
Company name: {data.company_name}
Hiring manager name: {data.hiring_manager_name}
Target location / country: {data.target_location}

Optional job context:
{data.job_context}

Personal background:
{data.personal_background}

Tasks:
1. Write a short LinkedIn connection request message.
2. Write a slightly longer direct outreach message.
3. Write a polite follow-up message if no reply is received.

Rules:
- Keep messages professional and concise.
- Do not use emojis.
- Do not sound desperate.
- Do not invent facts.
- Mention the role and company naturally.
- The follow-up should be polite and low-pressure.

Return ONLY valid JSON in this exact format:
{{
  "connection_message": "string",
  "outreach_message": "string",
  "follow_up_message": "string"
}}
"""

    resp = client.chat.completions.create(
        model="gpt-4.1-mini",
        messages=[
            {"role": "system", "content": system_msg},
            {"role": "user", "content": user_msg},
        ],
        temperature=0.3,
    )

    content = (resp.choices[0].message.content or "").strip()
    parsed = parse_json_response(content)

    required_keys = {"connection_message", "outreach_message", "follow_up_message"}
    if not required_keys.issubset(parsed.keys()):
        raise HTTPException(
            status_code=500,
            detail=f"Missing keys in message response. Required: {required_keys}. Got: {list(parsed.keys())}",
        )

    return parsed


@app.post("/save-job-alert")
def save_job_alert(data: JobAlertInput):
    save_json_file(JOB_ALERT_FILE, data.model_dump())
    return {"status": "saved", "message": "Job alert preferences saved successfully."}


@app.post("/send-test-job-alert", response_model=JobAlertTestOutput)
def send_test_job_alert(data: JobAlertInput):
    subject, body = build_job_alert_preview(data)
    save_json_file(JOB_ALERT_FILE, data.model_dump())

    if SMTP_HOST and SMTP_USER and SMTP_PASSWORD and SMTP_FROM:
        try:
            send_email(data.email, subject, body)
            return JobAlertTestOutput(
                status="sent",
                message="Test email sent successfully.",
                preview_subject=subject,
                preview_body=body,
            )
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Failed to send test email: {e}")

    return JobAlertTestOutput(
        status="preview_only",
        message="SMTP is not configured yet. Preview generated successfully.",
        preview_subject=subject,
        preview_body=body,
    )


@app.get("/job-alert-config")
def get_job_alert_config():
    data = load_json_file(JOB_ALERT_FILE)
    return {"status": "ok", "config": data}


@app.post("/format-linkedin")
def format_linkedin(data: LinkedInCopyInput):
    exp = "\n".join([f"• {b}" for b in data.experience_bullets])

    copy_text = f"""HEADLINE
{data.headline}

ABOUT
{data.about}

EXPERIENCE HIGHLIGHTS
{exp}
"""

    return {"copy_text": copy_text}
