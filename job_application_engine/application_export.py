from __future__ import annotations

from io import BytesIO
from pathlib import Path
import json
import re

from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import A4, letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


def _slug(value: str) -> str:
    slug = re.sub(r'[^a-z0-9]+', '-', str(value or '').strip().lower()).strip('-')
    return slug or 'job-application-report'


def _page_size_for_country(target_country: str):
    if (target_country or '').strip().lower() in {'usa', 'united states', 'canada'}:
        return letter
    return A4


def _render_docx(report: dict) -> BytesIO:
    document = Document()
    document.styles['Normal'].font.name = 'Calibri'
    document.styles['Normal'].font.size = Pt(11)
    document.add_heading('Job Application Intelligence Report', level=0)
    for heading, body in report.items():
        document.add_heading(heading.replace('_', ' ').title(), level=1)
        if isinstance(body, dict):
            for sub_heading, sub_body in body.items():
                document.add_paragraph(sub_heading.replace('_', ' ').title(), style='List Bullet')
                if isinstance(sub_body, list):
                    for item in sub_body:
                        document.add_paragraph(str(item), style='List Bullet 2')
                else:
                    document.add_paragraph(str(sub_body))
        elif isinstance(body, list):
            for item in body:
                document.add_paragraph(str(item), style='List Bullet')
        else:
            document.add_paragraph(str(body))
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def _render_pdf(report: dict, target_country: str) -> BytesIO:
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=_page_size_for_country(target_country), leftMargin=42, rightMargin=42, topMargin=42, bottomMargin=42)
    styles = getSampleStyleSheet()
    heading_style = ParagraphStyle('Heading', parent=styles['Heading2'], fontName='Helvetica-Bold', fontSize=11, leading=14, spaceBefore=8, spaceAfter=4)
    body_style = ParagraphStyle('Body', parent=styles['BodyText'], fontName='Helvetica', fontSize=10, leading=14, spaceAfter=4)
    story = [Paragraph('Job Application Intelligence Report', styles['Title']), Spacer(1, 10)]
    for heading, body in report.items():
        story.append(Paragraph(heading.replace('_', ' ').title(), heading_style))
        if isinstance(body, dict):
            for sub_heading, sub_body in body.items():
                story.append(Paragraph(f'<b>{sub_heading.replace("_", " ").title()}</b>', body_style))
                if isinstance(sub_body, list):
                    for item in sub_body:
                        story.append(Paragraph(f'&bull; {str(item)}', body_style))
                else:
                    story.append(Paragraph(str(sub_body).replace('\n', '<br/>'), body_style))
        elif isinstance(body, list):
            for item in body:
                story.append(Paragraph(f'&bull; {str(item)}', body_style))
        else:
            story.append(Paragraph(str(body).replace('\n', '<br/>'), body_style))
        story.append(Spacer(1, 4))
    doc.build(story)
    buffer.seek(0)
    return buffer


def export_application_report(full_name: str, target_role: str, target_country: str, report: dict) -> dict:
    out_dir = Path('rendered') / 'job-applications'
    out_dir.mkdir(parents=True, exist_ok=True)
    slug = _slug('-'.join(part for part in [full_name, target_role, 'application-report'] if str(part or '').strip()))
    docx_path = out_dir / f'{slug}.docx'
    pdf_path = out_dir / f'{slug}.pdf'
    json_path = out_dir / f'{slug}.json'
    docx_path.write_bytes(_render_docx(report).getvalue())
    pdf_path.write_bytes(_render_pdf(report, target_country=target_country).getvalue())
    json_path.write_text(json.dumps(report, indent=2), encoding='utf-8')
    return {
        'application_report_docx_path': str(docx_path.resolve()),
        'application_report_pdf_path': str(pdf_path.resolve()),
        'application_report_json_path': str(json_path.resolve()),
    }
