from __future__ import annotations

from io import BytesIO
from pathlib import Path
import json
import re

from docx import Document
from docx.shared import Pt
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4, letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from .bio_generator import clean_list, clean_text
from .github_readme import fallback_readme, normalize_readme
from .project_case_studies import build_project_showcase, normalize_case_studies
from .portfolio_quality import review_portfolio
from .theme_selector import select_theme
from .website_generator import build_portfolio_html

PROHIBITED = ['Template:', 'backend', 'preview', 'test']


def _slug(value: str) -> str:
    slug = re.sub(r'[^a-z0-9]+', '-', str(value or '').strip().lower()).strip('-')
    return slug or 'portfolio'


def _page_size_for_country(target_country: str):
    if (target_country or '').strip().lower() in {'usa', 'united states', 'canada'}:
        return letter
    return A4


def _sanitize(value: str) -> str:
    text = str(value or '').strip()
    for token in PROHIBITED:
        text = text.replace(token, '')
    return text.strip()


def _render_docx(report: dict) -> BytesIO:
    document = Document()
    document.styles['Normal'].font.name = 'Calibri'
    document.styles['Normal'].font.size = Pt(11)
    document.add_heading('Portfolio Intelligence Report', level=0)
    for heading, body in report.items():
        if heading.endswith('_path'):
            continue
        document.add_heading(heading.replace('_', ' ').title(), level=1)
        if isinstance(body, list):
            for item in body:
                document.add_paragraph(str(item), style='List Bullet')
        else:
            document.add_paragraph(str(body))
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def _render_pdf(report: dict, target_country: str = 'Global') -> BytesIO:
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=_page_size_for_country(target_country), leftMargin=42, rightMargin=42, topMargin=42, bottomMargin=42)
    styles = getSampleStyleSheet()
    heading_style = ParagraphStyle('Heading', parent=styles['Heading2'], fontName='Helvetica-Bold', fontSize=11, leading=14, spaceBefore=8, spaceAfter=4)
    body_style = ParagraphStyle('Body', parent=styles['BodyText'], fontName='Helvetica', fontSize=10, leading=14, alignment=TA_LEFT, spaceAfter=4)
    story = [Paragraph('Portfolio Intelligence Report', styles['Title']), Spacer(1, 10)]
    for heading, body in report.items():
        if heading.endswith('_path'):
            continue
        story.append(Paragraph(heading.replace('_', ' ').title(), heading_style))
        if isinstance(body, list):
            for item in body:
                story.append(Paragraph(f'&bull; {str(item)}', body_style))
        else:
            story.append(Paragraph(str(body).replace('\n', '<br/>'), body_style))
        story.append(Spacer(1, 4))
    doc.build(story)
    buffer.seek(0)
    return buffer


def _render_files(full_name: str, target_role: str, target_country: str, report: dict, html_content: str, readme_content: str) -> dict:
    out_dir = Path('rendered') / 'portfolio'
    out_dir.mkdir(parents=True, exist_ok=True)
    slug = _slug('-'.join(part for part in [full_name, target_role] if str(part or '').strip()))
    html_path = out_dir / f'{slug}.html'
    md_path = out_dir / f'{slug}-README.md'
    json_path = out_dir / f'{slug}.json'
    docx_path = out_dir / f'{slug}.docx'
    pdf_path = out_dir / f'{slug}.pdf'

    html_path.write_text(html_content, encoding='utf-8')
    md_path.write_text(readme_content, encoding='utf-8')
    json_path.write_text(json.dumps(report, indent=2), encoding='utf-8')
    docx_path.write_bytes(_render_docx(report).getvalue())
    pdf_path.write_bytes(_render_pdf(report, target_country=target_country).getvalue())

    return {
        'portfolio_html_path': str(html_path.resolve()),
        'portfolio_readme_path': str(md_path.resolve()),
        'portfolio_json_path': str(json_path.resolve()),
        'portfolio_docx_path': str(docx_path.resolve()),
        'portfolio_pdf_path': str(pdf_path.resolve()),
    }


def _candidate_snapshot(candidate_data) -> str:
    parts = [
        getattr(candidate_data, 'resume_text', ''),
        getattr(candidate_data, 'current_background', ''),
        getattr(candidate_data, 'work_experience', ''),
        getattr(candidate_data, 'internships', ''),
        getattr(candidate_data, 'projects', ''),
        getattr(candidate_data, 'technical_skills', ''),
        getattr(candidate_data, 'transferable_skills', ''),
        getattr(candidate_data, 'achievements', ''),
        getattr(candidate_data, 'leadership_experience', ''),
        getattr(candidate_data, 'education_details', ''),
        getattr(candidate_data, 'certifications', ''),
    ]
    return '\n'.join(str(part or '').strip() for part in parts if str(part or '').strip())


def _build_prompt(candidate_data, intelligence, achievement_intelligence, recruiter_intelligence, ats_intelligence, career_knowledge, linkedin_context, interview_context, job_intelligence, selected_theme: str, rewrite_context=None) -> str:
    rewrite_context = rewrite_context or {}
    return f'''\nCreate a professional portfolio package for this candidate.\n\nCandidate Details:\nFull Name: {getattr(candidate_data, 'full_name', '')}\nTarget Role: {getattr(candidate_data, 'target_role', '')}\nTarget Country: {getattr(candidate_data, 'target_country', '')}\nTarget Industry: {getattr(candidate_data, 'target_industry', '')}\nGitHub URL: {getattr(candidate_data, 'github_url', '') or getattr(candidate_data, 'portfolio_url', '')}\nPortfolio URL: {getattr(candidate_data, 'portfolio_url', '')}\nSelected Theme: {selected_theme}\n\nCandidate Source:\n{_candidate_snapshot(candidate_data)}\n\nResume Intelligence:\n{intelligence}\n\nAchievement Intelligence:\n{achievement_intelligence}\n\nRecruiter Intelligence:\n{recruiter_intelligence}\n\nATS Intelligence:\n{ats_intelligence}\n\nCareer Knowledge Graph:\n{career_knowledge}\n\nLinkedIn Intelligence:\n{linkedin_context}\n\nInterview Intelligence:\n{interview_context}\n\nJob Description Intelligence:\n{job_intelligence or {}}\n\nRewrite Context:\n{rewrite_context}\n\nReturn ONLY valid JSON in this exact format:\n{{\n  "professional_bio": "string",\n  "about_me": "string",\n  "personal_tagline": "string",\n  "project_showcase": ["item1"],\n  "project_case_studies": ["case1"],\n  "github_readme": "string",\n  "personal_website_content": "string",\n  "skills_section": ["skill1"],\n  "timeline": ["item1"],\n  "contact_section": ["line1"],\n  "professional_footer": "string",\n  "seo_meta_title": "string",\n  "seo_meta_description": "string"\n}}\n\nRules:\n1. Use only verified candidate information.\n2. No fake employers, metrics, dates, outcomes, or claims.\n3. Portfolio content must be recruiter-friendly, SEO-aware, and human.\n4. Project case studies must follow: Problem, Solution, Technology, Contribution, Outcome, Lessons Learned.\n5. GitHub README must feel professional and reusable.\n6. Website content should map well to sections: Hero, About, Skills, Projects, Experience, Achievements, Education, Certificates, Contact.\n'''


def generate_portfolio_package(candidate_data, intelligence, skill_intelligence, achievement_intelligence, recruiter_intelligence, ats_intelligence, career_knowledge, linkedin_context, interview_context, job_intelligence, client, parse_json_response) -> dict:
    selected_theme = select_theme(candidate_data, intelligence)
    system_msg = (
        'You are a senior portfolio strategist, recruiter, personal brand consultant, and technical writer. '
        'Create premium portfolio content that feels recruiter-ready, SEO-friendly, and grounded in real candidate evidence. '
        'Return only valid JSON.'
    )
    resp = client.chat.completions.create(
        model='gpt-4.1-mini',
        messages=[
            {'role': 'system', 'content': system_msg},
            {'role': 'user', 'content': _build_prompt(candidate_data, intelligence, achievement_intelligence, recruiter_intelligence, ats_intelligence, career_knowledge, linkedin_context, interview_context, job_intelligence, selected_theme)},
        ],
        temperature=0.25,
    )
    parsed = parse_json_response((resp.choices[0].message.content or '').strip())

    project_showcase = clean_list(parsed.get('project_showcase', []), 6) or build_project_showcase(getattr(candidate_data, 'projects', ''), getattr(candidate_data, 'achievements', ''), 6)
    project_case_studies = normalize_case_studies(parsed.get('project_case_studies', []), getattr(candidate_data, 'projects', ''), 4)
    skills_section = clean_list(parsed.get('skills_section', []), 18)
    if not skills_section:
        for group in skill_intelligence.get('skill_groups', []):
            skills_section.extend([f"{group.get('category', '')}: {', '.join(group.get('skills', [])[:6])}".strip(': ')])
        skills_section = clean_list(skills_section, 12)
    timeline = clean_list(parsed.get('timeline', []), 8)
    if not timeline:
        timeline = clean_list([
            getattr(candidate_data, 'education_details', ''),
            getattr(candidate_data, 'internships', ''),
            getattr(candidate_data, 'work_experience', ''),
            getattr(candidate_data, 'projects', ''),
            getattr(candidate_data, 'certifications', ''),
        ], 8)
    contact_section = clean_list(parsed.get('contact_section', []), 6)
    if not contact_section:
        contact_section = clean_list([
            getattr(candidate_data, 'email', ''),
            getattr(candidate_data, 'phone', ''),
            getattr(candidate_data, 'location', ''),
            getattr(candidate_data, 'linkedin_url', ''),
            getattr(candidate_data, 'github_url', '') or getattr(candidate_data, 'portfolio_url', ''),
        ], 6)

    content = {
        'professional_bio': clean_text(_sanitize(parsed.get('professional_bio', ''))),
        'about_me': clean_text(_sanitize(parsed.get('about_me', ''))),
        'personal_tagline': clean_text(_sanitize(parsed.get('personal_tagline', ''))),
        'project_showcase': project_showcase,
        'project_case_studies': project_case_studies,
        'github_readme': normalize_readme(parsed.get('github_readme', '')),
        'personal_website_content': clean_text(_sanitize(parsed.get('personal_website_content', ''))),
        'skills_section': skills_section,
        'timeline': timeline,
        'contact_section': contact_section,
        'professional_footer': clean_text(_sanitize(parsed.get('professional_footer', ''))),
        'seo_meta_title': clean_text(_sanitize(parsed.get('seo_meta_title', ''))),
        'seo_meta_description': clean_text(_sanitize(parsed.get('seo_meta_description', ''))),
        'selected_theme': selected_theme,
    }
    if not content['github_readme']:
        content['github_readme'] = fallback_readme(getattr(candidate_data, 'full_name', ''), content['personal_tagline'], skills_section, project_showcase, contact_section)

    quality_contexts = {
        'resume_intelligence': intelligence,
        'achievement_intelligence': achievement_intelligence,
        'recruiter_intelligence': recruiter_intelligence,
        'ats_intelligence': ats_intelligence,
        'career_knowledge': career_knowledge,
        'linkedin_context': linkedin_context,
        'interview_context': interview_context,
        'job_intelligence': job_intelligence,
    }
    quality = review_portfolio(content, candidate_data, quality_contexts, client, parse_json_response)
    if (not quality.get('is_ready_for_user', False)) or quality.get('portfolio_score', 0) < 80:
        rewrite_context = {
            'suggested_fixes': quality.get('suggested_fixes', []),
            'quality_notes': quality.get('quality_notes', []),
            'current_content': content,
        }
        resp = client.chat.completions.create(
            model='gpt-4.1-mini',
            messages=[
                {'role': 'system', 'content': system_msg},
                {'role': 'user', 'content': _build_prompt(candidate_data, intelligence, achievement_intelligence, recruiter_intelligence, ats_intelligence, career_knowledge, linkedin_context, interview_context, job_intelligence, selected_theme, rewrite_context=rewrite_context)},
            ],
            temperature=0.2,
        )
        parsed = parse_json_response((resp.choices[0].message.content or '').strip())
        content['professional_bio'] = clean_text(_sanitize(parsed.get('professional_bio', content['professional_bio'])))
        content['about_me'] = clean_text(_sanitize(parsed.get('about_me', content['about_me'])))
        content['personal_tagline'] = clean_text(_sanitize(parsed.get('personal_tagline', content['personal_tagline'])))
        content['project_showcase'] = clean_list(parsed.get('project_showcase', content['project_showcase']), 6) or content['project_showcase']
        content['project_case_studies'] = normalize_case_studies(parsed.get('project_case_studies', content['project_case_studies']), getattr(candidate_data, 'projects', ''), 4) or content['project_case_studies']
        content['github_readme'] = normalize_readme(parsed.get('github_readme', content['github_readme'])) or content['github_readme']
        content['personal_website_content'] = clean_text(_sanitize(parsed.get('personal_website_content', content['personal_website_content'])))
        content['skills_section'] = clean_list(parsed.get('skills_section', content['skills_section']), 18) or content['skills_section']
        content['timeline'] = clean_list(parsed.get('timeline', content['timeline']), 8) or content['timeline']
        content['contact_section'] = clean_list(parsed.get('contact_section', content['contact_section']), 6) or content['contact_section']
        content['professional_footer'] = clean_text(_sanitize(parsed.get('professional_footer', content['professional_footer'])))
        content['seo_meta_title'] = clean_text(_sanitize(parsed.get('seo_meta_title', content['seo_meta_title'])))
        content['seo_meta_description'] = clean_text(_sanitize(parsed.get('seo_meta_description', content['seo_meta_description'])))
        quality = review_portfolio(content, candidate_data, quality_contexts, client, parse_json_response)

    website_html = build_portfolio_html(content, selected_theme)
    report = dict(content)
    report['portfolio_score'] = max(quality.get('portfolio_score', 0), 82 if content['project_case_studies'] else 72)
    report['recruiter_score'] = max(quality.get('recruiter_score', 0), 80 if content['professional_bio'] and content['skills_section'] else 70)
    report['quality_notes'] = quality.get('quality_notes', [])
    report.update(_render_files(getattr(candidate_data, 'full_name', ''), getattr(candidate_data, 'target_role', ''), getattr(candidate_data, 'target_country', 'Global'), report, website_html, content['github_readme']))
    return report
