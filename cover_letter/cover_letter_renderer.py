from __future__ import annotations

from io import BytesIO
from pathlib import Path
import re

from docx import Document
from docx.shared import Pt
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4, letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


def _slug(value: str) -> str:
    slug = re.sub(r"[^a-z0-9]+", "-", (value or "").strip().lower()).strip("-")
    return slug or "cover-letter"


def _page_size_for_country(target_country: str):
    if (target_country or "").strip().lower() in {"usa", "united states", "canada"}:
        return letter
    return A4


def build_cover_letter_docx(cover_letter_text: str) -> BytesIO:
    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for block in [block.strip() for block in str(cover_letter_text or "").split("\n\n") if block.strip()]:
        paragraph = document.add_paragraph()
        for line in block.splitlines():
            paragraph.add_run(line.strip())
            paragraph.add_run("\n")
        if paragraph.runs:
            paragraph.runs[-1].text = paragraph.runs[-1].text.rstrip("\n")

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def build_cover_letter_pdf(cover_letter_text: str, target_country: str = "Global") -> BytesIO:
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=_page_size_for_country(target_country), leftMargin=54, rightMargin=54, topMargin=54, bottomMargin=54)
    styles = getSampleStyleSheet()
    body_style = ParagraphStyle(
        "CoverLetterBody",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=10.5,
        leading=15,
        alignment=TA_LEFT,
        spaceAfter=8,
    )
    story = []
    for block in [block.strip() for block in str(cover_letter_text or "").split("\n\n") if block.strip()]:
        safe_block = "<br/>".join(line.strip() for line in block.splitlines() if line.strip())
        story.append(Paragraph(safe_block, body_style))
        story.append(Spacer(1, 4))
    doc.build(story)
    buffer.seek(0)
    return buffer


def render_cover_letter_package(full_name: str, target_role: str, company_name: str, cover_letter_text: str, target_country: str = "Global", output_dir: str | None = None) -> dict:
    base_dir = Path(output_dir) if output_dir else Path("rendered") / "cover_letters"
    base_dir.mkdir(parents=True, exist_ok=True)
    slug = _slug("-".join(part for part in [full_name, company_name, target_role] if (part or "").strip()))
    docx_path = base_dir / f"{slug}.docx"
    pdf_path = base_dir / f"{slug}.pdf"

    docx_buffer = build_cover_letter_docx(cover_letter_text)
    pdf_buffer = build_cover_letter_pdf(cover_letter_text, target_country=target_country)

    docx_path.write_bytes(docx_buffer.getvalue())
    pdf_path.write_bytes(pdf_buffer.getvalue())

    return {
        "docx_buffer": docx_buffer,
        "pdf_buffer": pdf_buffer,
        "cover_letter_docx_path": str(docx_path.resolve()),
        "cover_letter_pdf_path": str(pdf_path.resolve()),
    }
