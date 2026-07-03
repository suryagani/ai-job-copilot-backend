from __future__ import annotations

from io import BytesIO
from pathlib import Path
import re

from docx import Document
from docx.shared import Pt
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4, letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from .behavioral_questions import fallback_behavioral_questions, normalize_behavioral_questions
from .company_research_engine import build_company_context
from .feedback_engine import normalize_feedback
from .hr_questions import fallback_hr_questions, normalize_hr_questions
from .mock_interview_engine import fallback_mock_plan, normalize_mock_plan
from .star_answer_engine import fallback_star_examples, normalize_star_examples
from .technical_questions import build_fallback_technical_questions, normalize_questions

PROHIBITED = ['Template:', 'backend', 'preview', 'test']


def _clean_list(values, limit: int | None = None) -> list[str]:
    cleaned = []
    for value in values or []:
        text = str(value or '').strip().lstrip('-? ')
        if text and text not in cleaned:
            cleaned.append(text)
    return cleaned[:limit] if limit else cleaned


def _sanitize(value: str) -> str:
    text = str(value or '').strip()
    for token in PROHIBITED:
        text = text.replace(token, '')
    return text.strip()


def _slug(value: str) -> str:
    slug = re.sub(r'[^a-z0-9]+', '-', str(value or '').strip().lower()).strip('-')
    return slug or 'interview-prep'


def _page_size_for_country(target_country: str):
    if (target_country or '').strip().lower() in {'usa', 'united states', 'canada'}:
        return letter
    return A4


def _render_docx(report: dict) -> BytesIO:
    document = Document()
    document.styles['Normal'].font.name = 'Calibri'
    document.styles['Normal'].font.size = Pt(11)
    document.add_heading('Interview Preparation Report', level=0)
    for heading, body in report.items():
        if heading in {'interview_report_pdf_path', 'interview_report_docx_path'}:
            continue
        document.add_heading(heading.replace('_', ' ').title(), level=1)
        if isinstance(body, list):
            for item in body:
                document.add_paragraph(str(item), style='List Bullet')
        else:
            document.add_paragraph(str(body))
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def _render_pdf(report: dict, target_country: str = 'Global') -> BytesIO:
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=_page_size_for_country(target_country), leftMargin=42, rightMargin=42, topMargin=42, bottomMargin=42)
    styles = getSampleStyleSheet()
    heading_style = ParagraphStyle('Heading', parent=styles['Heading2'], fontName='Helvetica-Bold', fontSize=11, leading=14, spaceBefore=8, spaceAfter=4)
    body_style = ParagraphStyle('Body', parent=styles['BodyText'], fontName='Helvetica', fontSize=10, leading=14, alignment=TA_LEFT, spaceAfter=4)
    story = [Paragraph('Interview Preparation Report', styles['Title']), Spacer(1, 10)]
    for heading, body in report.items():
        if heading in {'interview_report_pdf_path', 'interview_report_docx_path'}:
            continue
        story.append(Paragraph(heading.replace('_', ' ').title(), heading_style))
        if isinstance(body, list):
            for item in body:
                story.append(Paragraph(f'&bull; {str(item)}', body_style))
        else:
            story.append(Paragraph(str(body).replace('\n', '<br/>'), body_style))
        story.append(Spacer(1, 4))
    doc.build(story)
    buffer.seek(0)
    return buffer


def _render_report_files(full_name: str, target_role: str, target_country: str, report: dict) -> dict:
    out_dir = Path('rendered') / 'interview'
    out_dir.mkdir(parents=True, exist_ok=True)
    slug = _slug('-'.join(part for part in [full_name, target_role] if str(part or '').strip()))
    docx_path = out_dir / f'{slug}.docx'
    pdf_path = out_dir / f'{slug}.pdf'
    docx_buffer = _render_docx(report)
    pdf_buffer = _render_pdf(report, target_country=target_country)
    docx_path.write_bytes(docx_buffer.getvalue())
    pdf_path.write_bytes(pdf_buffer.getvalue())
    return {
        'interview_report_docx_path': str(docx_path.resolve()),
        'interview_report_pdf_path': str(pdf_path.resolve()),
    }


def _candidate_snapshot(candidate_data) -> str:
    parts = [
        getattr(candidate_data, 'resume_text', ''),
        getattr(candidate_data, 'current_background', ''),
        getattr(candidate_data, 'work_experience', ''),
        getattr(candidate_data, 'internships', ''),
        getattr(candidate_data, 'projects', ''),
        getattr(candidate_data, 'technical_skills', ''),
        getattr(candidate_data, 'transferable_skills', ''),
        getattr(candidate_data, 'achievements', ''),
        getattr(candidate_data, 'leadership_experience', ''),
    ]
    return '\n'.join(str(part or '').strip() for part in parts if str(part or '').strip())


def _build_prompt(candidate_data, intelligence, job_intelligence, recruiter_intelligence, ats_intelligence, career_knowledge, company_context, rewrite_context=None) -> str:
    rewrite_context = rewrite_context or {}
    return f'''\nCreate personalized interview preparation for this candidate.\n\nCandidate Details:\nFull Name: {getattr(candidate_data, 'full_name', '')}\nTarget Role: {getattr(candidate_data, 'target_role', '')}\nTarget Company: {getattr(candidate_data, 'company_name', '')}\nTarget Country: {getattr(candidate_data, 'target_country', '')}\nIndustry: {getattr(candidate_data, 'target_industry', '')}\nYears of Experience: {getattr(candidate_data, 'years_of_experience', '') or getattr(candidate_data, 'experience_level', '')}\n\nResume / Candidate Source:\n{_candidate_snapshot(candidate_data)}\n\nResume Intelligence:\n{intelligence}\n\nJob Description Intelligence:\n{job_intelligence or {}}\n\nRecruiter Intelligence:\n{recruiter_intelligence or {}}\n\nATS Intelligence:\n{ats_intelligence or {}}\n\nCareer Knowledge Graph:\n{career_knowledge or {}}\n\nCompany Context:\n{company_context}\n\nRewrite Context:\n{rewrite_context}\n\nReturn ONLY valid JSON in this exact format:\n{{\n  "technical_questions": ["q1"],\n  "behavioral_questions": ["q1"],\n  "hr_questions": ["q1"],\n  "company_specific_questions": ["q1"],\n  "star_answer_examples": ["example1"],\n  "candidate_strengths": ["strength1"],\n  "candidate_weaknesses": ["weakness1"],\n  "likely_follow_up_questions": ["q1"],\n  "interview_tips": ["tip1"],\n  "confidence_score": 0,\n  "readiness_score": 0,\n  "mock_interview_plan": ["step1"]\n}}\n\nRules:\n1. Questions must be role-specific and practical.\n2. If company is supplied, include company-aware questions and expectations.\n3. STAR answers must use only candidate information and frameworks, never fake experience.\n4. Do not invent metrics, employers, or achievements.\n5. Keep the output recruiter-level, realistic, and useful for actual interview preparation.\n'''


def generate_interview_prep_package(candidate_data, intelligence, job_intelligence, recruiter_intelligence, ats_intelligence, career_knowledge, client, parse_json_response) -> dict:
    company_context = build_company_context(
        company_name=getattr(candidate_data, 'company_name', ''),
        target_role=getattr(candidate_data, 'target_role', ''),
        target_country=getattr(candidate_data, 'target_country', 'Global'),
        industry=getattr(candidate_data, 'target_industry', ''),
        job_intelligence=job_intelligence,
    )

    system_msg = (
        'You are a senior interviewer, recruiter, hiring manager, and interview coach with 15 years of experience. '
        'Create tailored interview preparation that feels specific to the candidate, role, and company. '
        'Never invent experience. Return only valid JSON.'
    )
    resp = client.chat.completions.create(
        model='gpt-4.1-mini',
        messages=[
            {'role': 'system', 'content': system_msg},
            {'role': 'user', 'content': _build_prompt(candidate_data, intelligence, job_intelligence, recruiter_intelligence, ats_intelligence, career_knowledge, company_context)},
        ],
        temperature=0.25,
    )
    parsed = parse_json_response((resp.choices[0].message.content or '').strip())

    technical = normalize_questions(parsed.get('technical_questions', []), build_fallback_technical_questions(getattr(candidate_data, 'target_role', ''), (job_intelligence or {}).get('required_skills', []), getattr(candidate_data, 'projects', '')), 8)
    behavioral = normalize_behavioral_questions(parsed.get('behavioral_questions', []), 8) or fallback_behavioral_questions(getattr(candidate_data, 'target_role', ''))
    hr = normalize_hr_questions(parsed.get('hr_questions', []), 8) or fallback_hr_questions(getattr(candidate_data, 'target_role', ''), getattr(candidate_data, 'company_name', ''))
    company_questions = normalize_questions(parsed.get('company_specific_questions', []), [
        f"Why do you want to join {company_context['company_name']} for this {getattr(candidate_data, 'target_role', '')} role?",
        f"How would you align your strengths with {', '.join(company_context['interview_expectations'][:3])}?",
    ], 6)
    star_examples = normalize_star_examples(parsed.get('star_answer_examples', []), 5) or fallback_star_examples(candidate_data)
    mock_plan = normalize_mock_plan(parsed.get('mock_interview_plan', []), 7) or fallback_mock_plan(getattr(candidate_data, 'target_role', ''))

    feedback = normalize_feedback(parsed, candidate_data, recruiter_intelligence, ats_intelligence)

    coverage_confidence = 55 + min(len(technical) * 3, 18) + min(len(behavioral) * 2, 10) + min(len(company_questions) * 2, 8)
    coverage_readiness = 52 + min(len(technical) * 3, 18) + min(len(mock_plan) * 2, 10) + min(len(star_examples) * 3, 12)
    result = {
        'technical_questions': technical,
        'behavioral_questions': behavioral,
        'hr_questions': hr,
        'company_specific_questions': company_questions,
        'star_answer_examples': star_examples,
        'candidate_strengths': feedback['candidate_strengths'],
        'candidate_weaknesses': feedback['candidate_weaknesses'],
        'likely_follow_up_questions': _clean_list(parsed.get('likely_follow_up_questions', []), 8),
        'interview_tips': feedback['interview_tips'],
        'confidence_score': max(feedback['confidence_score'], min(100, coverage_confidence)),
        'readiness_score': max(feedback['readiness_score'], min(100, coverage_readiness)),
        'mock_interview_plan': mock_plan,
    }

    if result['readiness_score'] < 75 or len(result['technical_questions']) < 4:
        rewrite_context = {
            'current_result': result,
            'company_expectations': company_context,
            'missing_emphasis': recruiter_intelligence or {},
        }
        resp = client.chat.completions.create(
            model='gpt-4.1-mini',
            messages=[
                {'role': 'system', 'content': system_msg},
                {'role': 'user', 'content': _build_prompt(candidate_data, intelligence, job_intelligence, recruiter_intelligence, ats_intelligence, career_knowledge, company_context, rewrite_context=rewrite_context)},
            ],
            temperature=0.2,
        )
        parsed = parse_json_response((resp.choices[0].message.content or '').strip())
        result['technical_questions'] = normalize_questions(parsed.get('technical_questions', []), result['technical_questions'], 8)
        result['behavioral_questions'] = normalize_behavioral_questions(parsed.get('behavioral_questions', []), 8) or result['behavioral_questions']
        result['hr_questions'] = normalize_hr_questions(parsed.get('hr_questions', []), 8) or result['hr_questions']
        result['company_specific_questions'] = normalize_questions(parsed.get('company_specific_questions', []), result['company_specific_questions'], 6)
        result['star_answer_examples'] = normalize_star_examples(parsed.get('star_answer_examples', []), 5) or result['star_answer_examples']
        updated_feedback = normalize_feedback(parsed, candidate_data, recruiter_intelligence, ats_intelligence)
        result['candidate_strengths'] = updated_feedback['candidate_strengths'] or result['candidate_strengths']
        result['candidate_weaknesses'] = updated_feedback['candidate_weaknesses'] or result['candidate_weaknesses']
        result['likely_follow_up_questions'] = _clean_list(parsed.get('likely_follow_up_questions', []), 8) or result['likely_follow_up_questions']
        result['interview_tips'] = updated_feedback['interview_tips'] or result['interview_tips']
        result['confidence_score'] = max(result['confidence_score'], updated_feedback['confidence_score'])
        result['readiness_score'] = max(result['readiness_score'], updated_feedback['readiness_score'])
        result['mock_interview_plan'] = normalize_mock_plan(parsed.get('mock_interview_plan', []), 7) or result['mock_interview_plan']

    result.update(_render_report_files(getattr(candidate_data, 'full_name', ''), getattr(candidate_data, 'target_role', ''), getattr(candidate_data, 'target_country', 'Global'), result))
    return result
