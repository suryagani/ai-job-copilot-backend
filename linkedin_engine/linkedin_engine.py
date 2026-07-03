from __future__ import annotations

from io import BytesIO
from pathlib import Path
import re

from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import A4, letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from .about_engine import clean_about
from .creator_engine import build_creator_profile_suggestions, build_networking_suggestions
from .experience_engine import normalize_experience_rewrite
from .featured_engine import recommend_featured_section
from .headline_engine import clean_headline, score_headline
from .keyword_engine import build_top_keywords
from .linkedin_quality import review_linkedin_profile
from .skills_engine import reorder_skills

PROHIBITED = ["Template:", "backend", "preview", "test"]


def _clean_list(values) -> list[str]:
    cleaned = []
    for value in values or []:
        text = str(value or "").strip().lstrip("-? ")
        if text and text not in cleaned:
            cleaned.append(text)
    return cleaned


def _sanitize_text(value: str) -> str:
    text = str(value or "").strip()
    for token in PROHIBITED:
        text = text.replace(token, "")
    return text.strip()


def _slug(value: str) -> str:
    slug = re.sub(r"[^a-z0-9]+", "-", str(value or "").strip().lower()).strip("-")
    return slug or "linkedin-report"


def _page_size_for_country(target_country: str):
    if (target_country or "").strip().lower() in {"usa", "united states", "canada"}:
        return letter
    return A4


def _build_resume_context(candidate_data) -> str:
    parts = [
        getattr(candidate_data, "resume_text", ""),
        getattr(candidate_data, "current_background", ""),
        getattr(candidate_data, "work_experience", ""),
        getattr(candidate_data, "internships", ""),
        getattr(candidate_data, "projects", ""),
        getattr(candidate_data, "achievements", ""),
        getattr(candidate_data, "leadership_experience", ""),
    ]
    return "\n".join(str(part or "").strip() for part in parts if str(part or "").strip())


def _render_linkedin_report_docx(report: dict) -> BytesIO:
    document = Document()
    document.styles["Normal"].font.name = "Calibri"
    document.styles["Normal"].font.size = Pt(11)
    document.add_heading("LinkedIn Optimization Report", level=0)
    for heading, body in report.items():
        if heading in {"linkedin_report_docx_path", "linkedin_report_pdf_path"}:
            continue
        document.add_heading(heading.replace("_", " ").title(), level=1)
        if isinstance(body, list):
            for item in body:
                document.add_paragraph(str(item), style="List Bullet")
        else:
            document.add_paragraph(str(body))
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def _render_linkedin_report_pdf(report: dict, target_country: str = "Global") -> BytesIO:
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=_page_size_for_country(target_country), leftMargin=42, rightMargin=42, topMargin=42, bottomMargin=42)
    styles = getSampleStyleSheet()
    heading_style = ParagraphStyle("Heading", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=11, leading=14, spaceBefore=8, spaceAfter=4)
    body_style = ParagraphStyle("Body", parent=styles["BodyText"], fontName="Helvetica", fontSize=10, leading=14, spaceAfter=4)
    story = [Paragraph("LinkedIn Optimization Report", styles["Title"]), Spacer(1, 10)]
    for heading, body in report.items():
        if heading in {"linkedin_report_docx_path", "linkedin_report_pdf_path"}:
            continue
        story.append(Paragraph(heading.replace("_", " ").title(), heading_style))
        if isinstance(body, list):
            for item in body:
                story.append(Paragraph(f"&bull; {str(item)}", body_style))
        else:
            story.append(Paragraph(str(body).replace("\n", "<br/>"), body_style))
        story.append(Spacer(1, 4))
    doc.build(story)
    buffer.seek(0)
    return buffer


def _estimate_visibility_score(output: dict) -> int:
    score = 35
    if str(output.get("professional_headline", "")).strip():
        score += 15
    if len(str(output.get("about_section", "")).split()) >= 180:
        score += 15
    if len(output.get("experience_rewrite", [])) >= 3:
        score += 10
    score += min(len(output.get("top_50_recruiter_keywords", [])), 20)
    score += min(len(output.get("skills_order", [])) // 2, 10)
    return max(0, min(100, score))


def _render_report_files(full_name: str, target_role: str, target_country: str, report: dict) -> dict:
    out_dir = Path("rendered") / "linkedin"
    out_dir.mkdir(parents=True, exist_ok=True)
    slug = _slug("-".join(part for part in [full_name, target_role] if str(part or "").strip()))
    docx_path = out_dir / f"{slug}.docx"
    pdf_path = out_dir / f"{slug}.pdf"
    docx_buffer = _render_linkedin_report_docx(report)
    pdf_buffer = _render_linkedin_report_pdf(report, target_country=target_country)
    docx_path.write_bytes(docx_buffer.getvalue())
    pdf_path.write_bytes(pdf_buffer.getvalue())
    return {
        "linkedin_report_docx_path": str(docx_path.resolve()),
        "linkedin_report_pdf_path": str(pdf_path.resolve()),
    }


def _build_prompt(candidate_data, intelligence, job_intelligence, recruiter_intelligence, ats_intelligence, personalization, career_knowledge, top_keywords, skill_order, rewrite_context=None) -> str:
    rewrite_context = rewrite_context or {}
    return f'''\nCreate a complete LinkedIn optimization report for this candidate.\n\nCandidate Details:\nFull Name: {getattr(candidate_data, "full_name", "")}\nTarget Role: {getattr(candidate_data, "target_role", "")}\nTarget Country: {getattr(candidate_data, "target_country", "")}\nTarget Industry: {getattr(candidate_data, "target_industry", "")}\nExperience Level: {getattr(candidate_data, "experience_level", "")}\nCurrent LinkedIn: {getattr(candidate_data, "current_linkedin", "")}\nCurrent Headline: {getattr(candidate_data, "current_headline", "")}\nCurrent About: {getattr(candidate_data, "current_about", "")}\n\nResume / Background Source:\n{_build_resume_context(candidate_data)}\n\nResume Intelligence:\n{intelligence}\n\nJob Intelligence:\n{job_intelligence or {}}\n\nRecruiter Intelligence:\n{recruiter_intelligence or {}}\n\nATS Intelligence:\n{ats_intelligence or {}}\n\nPersonalization:\n{personalization or {}}\n\nCareer Knowledge Graph:\n{career_knowledge or {}}\n\nPriority Keywords:\n{top_keywords}\n\nSuggested Skills Order:\n{skill_order}\n\nRewrite Context:\n{rewrite_context}\n\nReturn ONLY valid JSON in this exact format:\n{{\n  "professional_headline": "string",\n  "about_section": "string",\n  "experience_rewrite": ["bullet1", "bullet2"],\n  "creator_profile_suggestions": ["suggestion1", "suggestion2"],\n  "networking_suggestions": ["suggestion1", "suggestion2"]\n}}\n\nRules:\n1. Headline maximum 220 characters.\n2. Headline must include role, specialization, value proposition, and recruiter-facing keywords naturally.\n3. About section must feel human, professional, and keyword-aware, roughly 300 to 500 words.\n4. Experience rewrite must be achievement-oriented, but never invent metrics, employers, dates, or responsibilities.\n5. Use only verified candidate information.\n6. No AI-sounding text, no emojis, no placeholder labels.\n7. Optimize the entire LinkedIn profile for recruiter discoverability.\n'''


def generate_linkedin_optimization_package(candidate_data, intelligence, skill_intelligence, job_intelligence, recruiter_intelligence, ats_intelligence, personalization, career_knowledge, client, parse_json_response) -> dict:
    top_keywords = build_top_keywords(skill_intelligence, ats_intelligence, job_intelligence, getattr(candidate_data, "target_role", ""), career_knowledge)
    skill_order = reorder_skills(skill_intelligence, top_keywords, technical_skills=getattr(candidate_data, "technical_skills", ""), transferable_skills=getattr(candidate_data, "transferable_skills", ""), tools_software=getattr(candidate_data, "tools_software", ""))

    system_msg = (
        "You are a senior LinkedIn branding strategist, recruiter, and executive profile writer. "
        "Write like a premium LinkedIn expert, not like generic AI. "
        "Optimize for recruiter visibility, authenticity, and professional positioning. Return only JSON."
    )
    user_msg = _build_prompt(candidate_data, intelligence, job_intelligence, recruiter_intelligence, ats_intelligence, personalization, career_knowledge, top_keywords, skill_order)
    resp = client.chat.completions.create(
        model="gpt-4.1-mini",
        messages=[
            {"role": "system", "content": system_msg},
            {"role": "user", "content": user_msg},
        ],
        temperature=0.3,
    )
    parsed = parse_json_response((resp.choices[0].message.content or "").strip())

    output = {
        "professional_headline": clean_headline(parsed.get("professional_headline", ""), getattr(candidate_data, "target_role", ""), top_keywords),
        "about_section": clean_about(_sanitize_text(parsed.get("about_section", ""))),
        "experience_rewrite": normalize_experience_rewrite(parsed.get("experience_rewrite", [])),
        "featured_section": recommend_featured_section(candidate_data, None, career_knowledge),
        "skills_order": skill_order,
        "top_50_recruiter_keywords": top_keywords,
        "creator_profile_suggestions": _clean_list(parsed.get("creator_profile_suggestions", [])) or build_creator_profile_suggestions(candidate_data, career_knowledge, recruiter_intelligence),
        "networking_suggestions": _clean_list(parsed.get("networking_suggestions", [])) or build_networking_suggestions(candidate_data, job_intelligence, career_knowledge),
    }

    quality_contexts = {
        "resume_intelligence": intelligence,
        "job_intelligence": job_intelligence,
        "recruiter_intelligence": recruiter_intelligence,
        "ats_intelligence": ats_intelligence,
        "personalization": personalization,
        "career_knowledge": career_knowledge,
    }
    quality = review_linkedin_profile(output, candidate_data, quality_contexts, client, parse_json_response)
    if (not quality.get("is_ready_for_user", False)) or quality.get("linkedin_score", 0) < 80:
        rewrite_context = {
            "suggested_fixes": quality.get("suggested_fixes", []),
            "quality_notes": quality.get("quality_notes", []),
            "previous_output": output,
        }
        resp = client.chat.completions.create(
            model="gpt-4.1-mini",
            messages=[
                {"role": "system", "content": system_msg},
                {"role": "user", "content": _build_prompt(candidate_data, intelligence, job_intelligence, recruiter_intelligence, ats_intelligence, personalization, career_knowledge, top_keywords, skill_order, rewrite_context=rewrite_context)},
            ],
            temperature=0.25,
        )
        parsed = parse_json_response((resp.choices[0].message.content or "").strip())
        output["professional_headline"] = clean_headline(parsed.get("professional_headline", ""), getattr(candidate_data, "target_role", ""), top_keywords)
        output["about_section"] = clean_about(_sanitize_text(parsed.get("about_section", "")))
        output["experience_rewrite"] = normalize_experience_rewrite(parsed.get("experience_rewrite", []))
        output["creator_profile_suggestions"] = _clean_list(parsed.get("creator_profile_suggestions", [])) or output["creator_profile_suggestions"]
        output["networking_suggestions"] = _clean_list(parsed.get("networking_suggestions", [])) or output["networking_suggestions"]
        quality = review_linkedin_profile(output, candidate_data, quality_contexts, client, parse_json_response)

    output["headline_score"] = score_headline(output["professional_headline"], getattr(candidate_data, "target_role", ""), top_keywords)
    output["linkedin_score"] = max(output["headline_score"], quality.get("linkedin_score", 0)) if quality.get("linkedin_score", 0) < output["headline_score"] else quality.get("linkedin_score", 0)
    output["recruiter_visibility_score"] = max(_estimate_visibility_score(output), quality.get("recruiter_visibility_score", 0))
    output["quality_notes"] = quality.get("quality_notes", [])
    output["visibility_explanation"] = quality.get("visibility_explanation", "")
    output.update(_render_report_files(getattr(candidate_data, "full_name", ""), getattr(candidate_data, "target_role", ""), getattr(candidate_data, "target_country", "Global"), output))
    return output
